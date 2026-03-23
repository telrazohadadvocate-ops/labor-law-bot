"""
Full End-to-End QA Test — Israeli Labor Law Case: ישראל ישראלי
Runs 10 tests against the labor-law-bot Flask application.
"""

import json
import os
import sys
import io
import sqlite3
import traceback
from datetime import date, datetime, timedelta

# Force UTF-8 output on Windows
if sys.stdout.encoding and sys.stdout.encoding.lower() not in ("utf-8", "utf8"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# Set DB_PATH BEFORE importing app so _init_db() uses the test DB
_TEST_DB = os.path.join(os.path.dirname(os.path.abspath(__file__)), "qa_test_cases.db")
os.environ["DB_PATH"] = _TEST_DB

from app import (
    app, calculate_all_claims, generate_claim_text,
    safe_float, safe_int, _compute_limitation, RECUPERATION_DAY_VALUE,
    KANBAN_STAGES, _init_db, DB_PATH,
)

# ── Colour output ────────────────────────────────────────────────────────────
GREEN = "\033[92m"
RED   = "\033[91m"
YELLOW= "\033[93m"
RESET = "\033[0m"
BOLD  = "\033[1m"

PASS_COUNT = 0
FAIL_COUNT = 0
RESULTS = []

def _ok(label, detail=""):
    global PASS_COUNT
    PASS_COUNT += 1
    RESULTS.append(("PASS", label))
    msg = f"  {GREEN}[PASS]{RESET} {label}"
    if detail:
        msg += f" — {detail}"
    print(msg)

def _fail(label, detail=""):
    global FAIL_COUNT
    FAIL_COUNT += 1
    RESULTS.append(("FAIL", label, detail))
    msg = f"  {RED}[FAIL]{RESET} {label}"
    if detail:
        msg += f"\n        {YELLOW}→ {detail}{RESET}"
    print(msg)

def section(title):
    print(f"\n{BOLD}{'='*60}{RESET}")
    print(f"{BOLD}  {title}{RESET}")
    print(f"{BOLD}{'='*60}{RESET}")

PASSWORD = "LT2026"

# ── Test case: ישראל ישראלי ────────────────────────────────────────────────
#
#  Employment: 01.03.2021 → 15.01.2026 (fired, no hearing)
#  Salary: 18,000 NIS gross / month
#  5 days/week, 9 h/day
#  Section 14 at 6% (not 8.33%)
#  Pension gap: first 6 months unpaid (zero deposited in those months)
#  Overtime: 10 h/week (all at 125% — 2h/day × 5 days)
#  Salary paid late 3 times
#  Vacation: 12 unused/unpaid days at termination
#  Recuperation: last year not paid (7 days)

TEST_DATA = {
    # Plaintiff
    "plaintiff_name":   "ישראל ישראלי",
    "plaintiff_id":     "012345678",
    "plaintiff_address":"רחוב הרצל 15 תל אביב",
    "plaintiff_phone":  "050-1234567",
    "gender":           "male",
    # Defendant
    "defendant_name":   "טכנולוגיה בע\"מ",
    "defendant_id":     "514567890",
    "defendant_type":   "company",
    "defendant_owner":  "דוד כהן",
    "defendant_business":"פיתוח תוכנה",
    "defendant_address":"רחוב רוטשילד 45 תל אביב",
    # Employment
    "start_date":       "2021-03-01",
    "end_date":         "2026-01-15",
    "job_title":        "מפתח תוכנה",
    "termination_type": "fired",
    "work_days_per_week":"5",
    "hours_per_day":    "9",
    "base_salary":      "18000",
    "commissions":      "0",
    "work_schedule":    "ימים א׳–ה׳, 9 שעות ביום",
    # Claims
    "claim_severance":    True,
    "claim_prior_notice": True,
    "claim_overtime":     True,
    "claim_pension":      True,
    "claim_vacation":     True,
    "claim_recuperation": True,
    "claim_salary_delay": True,
    # Severance — section 14 at 6% (employer deposited nothing directly)
    "severance_deposited":  "0",
    "section14_arrangement":True,
    "section14_rate":       "6",
    # Overtime — 10 h/week all at 125% (2h/day × 5 days)
    "weekly_overtime_125":  "10",
    "weekly_overtime_150":  "0",
    # Pension — 6 months were not deposited at all
    # Full 58 months owed; only 52 months deposited → gap = 6 × 18000 × 6.5%
    "pension_deposited":    str(round(52 * 18000 * 0.065, 2)),  # 60840
    # Vacation — 12 unused days not paid at termination
    "vacation_days_paid":   "0",
    "vacation_rate_paid":   "0",
    # Recuperation — last year (7 days) not paid
    "recuperation_days_paid":"0",
    # Salary delay — 3 late payments × ~5% penalty on 18000
    "salary_delay_amount":  "2700",
    # Court / attorney
    "court_header":  "בית הדין האזורי לעבודה בתל אביב",
    "attorney_name": "דורון לוין ו/או אוהד תלרז",
    "attorney_id":   "87732",
    "firm_name":     "לוין תלרז, משרד עורכי דין",
    "firm_address":  "מרחוב הארבעה 28, תל אביב, מגדל דרומי קומה 19",
    "firm_phone":    "054-3699782",
    "firm_fax":      "076-5102966",
    "firm_email":    "law@levin-telraz.co.il",
    # Raw narrative for AI
    "raw_text": (
        "פוטרתי ללא שימוע לאחר כ-5 שנות עבודה. "
        "במהלך עבודתי עבדתי שעות נוספות רבות ללא תגמול, "
        "הפנסיה הופרשה באיחור של חצי שנה ובשיעור חסר, "
        "לא קיבלתי דמי הבראה בשנה האחרונה, "
        "והשכר שולם באיחור מספר פעימות. "
        "לא קיבלתי מכתב פיטורים רשמי ולא שולמו לי פיצויי פיטורים עד היום."
    ),
    "narrative": (
        "פוטרתי ללא שימוע לאחר כ-5 שנות עבודה. "
        "במהלך עבודתי עבדתי שעות נוספות רבות ללא תגמול, "
        "הפנסיה הופרשה באיחור של חצי שנה ובשיעור חסר, "
        "לא קיבלתי דמי הבראה בשנה האחרונה, "
        "והשכר שולם באיחור מספר פעימות. "
        "לא קיבלתי מכתב פיטורים רשמי ולא שולמו לי פיצויי פיטורים עד היום."
    ),
}

# ─────────────────────────────────────────────────────────────────────────────
# Expected calculation values (derived from system logic, verified manually)
# ─────────────────────────────────────────────────────────────────────────────
# Duration: 2021-03-01 → 2026-01-15
#   relativedelta → 4y 10m 14d → total_months=58, decimal_years=4.83
EXPECTED_MONTHS        = 58
EXPECTED_DEC_YEARS     = 4.83
EXPECTED_SALARY        = 18000.0
# hourly = 18000 / (5 × 9 × 4.33) = 18000/194.85 ≈ 92.38
EXPECTED_HOURLY        = round(18000 / (5 * 9 * 4.33), 2)   # 92.38
# daily  = 18000 / (5 × 4.33) = 18000/21.65 ≈ 831.41
EXPECTED_DAILY         = round(18000 / (5 * 4.33), 2)        # 831.41

# Severance with Section 14 @ 6%:
#   statutory   = 18000 × 4.83 = 86,940
#   s14_contrib = 6/100 × 18000 × 58 = 62,640
#   net         = max(0, 86940 − 62640) = 24,300
EXPECTED_SEVERANCE_STATUTORY = round(18000 * EXPECTED_DEC_YEARS, 2)  # 86940
EXPECTED_S14_CONTRIB         = round(6/100 * 18000 * EXPECTED_MONTHS, 2)  # 62640
EXPECTED_SEVERANCE_NET       = max(0, round(EXPECTED_SEVERANCE_STATUTORY - EXPECTED_S14_CONTRIB, 2))  # 24300

# Prior notice: >12 months → 30 days
#   daily_rate_notice = 18000/30 = 600
#   amount = 30 × 600 = 18,000
EXPECTED_NOTICE        = 18000.0

# Vacation: system computes total entitled days, subtracts paid_days (0 here)
#   VACATION_DAYS_5DAY: y1=12,y2=12,y3=12,y4=12 → 48 + frac(y5)*13
#   frac = 4.83 − 4 = 0.83 → 0.83×13=10.79 → total≈58.79 days
#   value = 58.79 × 831.41 = 48,879 (approx)
EXPECTED_VAC_TOTAL_DAYS = round(48 + 0.83 * 13, 2)   # 58.79

# Recuperation for 4.83 years (all at 418.18/day, 0 paid):
#   y1=5, y2=6, y3=6, y4=7 → 24 + 0.83×7=5.81 → 29.81 days
#   value = 29.81 × 418.18 = 12,465.95
EXPECTED_REC_DAYS      = round(24 + 0.83 * 7, 2)   # 29.81
EXPECTED_REC_VALUE     = round(EXPECTED_REC_DAYS * RECUPERATION_DAY_VALUE, 2)  # 12465.95

# Pension gap (6 months missing at 6.5%):
#   total_owed = 58 × 18000 × 6.5% = 67,860
#   deposited  = 52 × 18000 × 6.5% = 60,840
#   gap        = 7,020
EXPECTED_PENSION_OWED       = round(18000 * EXPECTED_MONTHS * 0.065, 2)   # 67860
EXPECTED_PENSION_DEPOSITED  = round(52 * 18000 * 0.065, 2)                # 60840
EXPECTED_PENSION_GAP        = round(EXPECTED_PENSION_OWED - EXPECTED_PENSION_DEPOSITED, 2)  # 7020

# Overtime (10 h/wk @ 125%):
#   surcharge_125 = hourly × 0.25 = 92.38 × 0.25 = 23.095
#   monthly = 10 × 4 × 23.095 = 923.8
#   total = 923.8 × 58 months ≈ 53,580.4
EXPECTED_OT_TOTAL = round((10 * 4 * round(EXPECTED_HOURLY * 0.25, 2)) * EXPECTED_MONTHS, 2)

# Salary delay damages (entered by user): 2,700
EXPECTED_SALARY_DELAY = 2700.0

TODAY = date.today()


# ─────────────────────────────────────────────────────────────────────────────
# TEST 1 — MANUAL FLOW via /generate-ai route
# ─────────────────────────────────────────────────────────────────────────────
def test1_generate_ai():
    section("TEST 1 — Manual Flow via /generate-ai")
    with app.test_client() as c:
        c.post("/login", data={"password": PASSWORD})
        resp = c.post(
            "/generate-ai",
            data=json.dumps(TEST_DATA),
            content_type="application/json",
        )
        if resp.status_code != 200:
            _fail("generate-ai HTTP 200", f"got {resp.status_code}")
            return
        result = resp.get_json()
        if not result:
            _fail("generate-ai returns JSON", "no JSON body")
            return
        _ok("generate-ai HTTP 200")

        if result.get("success"):
            _ok("generate-ai success=True")
        else:
            _fail("generate-ai success", f"error: {result.get('error')}")
            return

        mode = result.get("mode", "")
        _ok(f"generate-ai mode reported", f"mode='{mode}'")

        claim_text = result.get("claim_text", "") or result.get("ai_text", "")
        if not claim_text:
            _fail("generate-ai returns claim_text", "claim_text is empty")
            return
        _ok("generate-ai returns claim_text", f"{len(claim_text)} chars")

        # Print first 300 chars of response
        print(f"\n    [AI RESPONSE PREVIEW — first 300 chars]:")
        print(f"    {repr(claim_text[:300])}\n")

        # Check: text is Hebrew (contains common Hebrew chars)
        hebrew_chars = sum(1 for ch in claim_text if '\u05d0' <= ch <= '\u05ea')
        if hebrew_chars > 50:
            _ok("claim_text is Hebrew", f"{hebrew_chars} Hebrew chars")
        else:
            _fail("claim_text should be Hebrew", f"only {hebrew_chars} Hebrew chars")

        # Check for === section delimiters (present in AI mode; template mode uses \n\n headers)
        if mode == "ai" and "===" in claim_text:
            _ok("AI mode: === section delimiters found")
        elif mode != "ai":
            # Template/fallback: check for section titles without ===
            if "כללי" in claim_text or "הצדדים" in claim_text or "תביעה" in claim_text:
                _ok("Template mode: section headers present in Hebrew")
            else:
                _fail("Template mode: expected Hebrew section headers")
        else:
            _fail("AI mode: no === section delimiters found")

        # Verify plaintiff name appears in text
        if "ישראל ישראלי" in claim_text or "ישראלי" in claim_text:
            _ok("Plaintiff name appears in claim text")
        else:
            _fail("Plaintiff name missing from claim text")

        # Verify no JSON artifacts
        json_artifacts = ["{", "}", '"claims"', '"amount":', '"success"']
        found_artifacts = [a for a in json_artifacts if a in claim_text]
        if not found_artifacts:
            _ok("No JSON artifacts in claim text")
        else:
            _fail("JSON artifacts found in claim text", f"found: {found_artifacts[:3]}")

        # Calculations must be present
        if result.get("calculations"):
            _ok("Calculations included in response")
        else:
            _fail("Calculations missing from generate-ai response")


# ─────────────────────────────────────────────────────────────────────────────
# TEST 2 — DOCX GENERATION
# ─────────────────────────────────────────────────────────────────────────────
def test2_docx_generation():
    section("TEST 2 — DOCX Generation")
    with app.test_client() as c:
        c.post("/login", data={"password": PASSWORD})

        # Use template mode (no AI key) — generate claim text first
        calc_resp = c.post("/calculate",
                           data=json.dumps(TEST_DATA),
                           content_type="application/json")
        calc_result = calc_resp.get_json()
        if not calc_result or not calc_result.get("success"):
            _fail("Pre-calc for DOCX", f"calc failed: {calc_result}")
            return
        _ok("Pre-calculation for DOCX succeeded")

        # Generate DOCX via /generate-docx (template mode)
        docx_data = dict(TEST_DATA)
        resp = c.post("/generate-docx",
                      data=json.dumps(docx_data),
                      content_type="application/json")

        if resp.status_code == 200:
            _ok("generate-docx HTTP 200")
        else:
            err = resp.get_json()
            _fail("generate-docx HTTP 200", f"got {resp.status_code}: {err}")
            return

        # Check content type
        ct = resp.content_type
        if "wordprocessingml" in ct or "octet-stream" in ct:
            _ok("DOCX content-type correct", ct)
        else:
            _fail("DOCX content-type", f"got {ct}")

        # Check Content-Disposition filename
        cd = resp.headers.get("Content-Disposition", "")
        if "ישראל ישראלי" in cd or "claim" in cd or ".docx" in cd:
            _ok("DOCX filename in Content-Disposition", cd[:80])
        else:
            _fail("DOCX filename", f"Content-Disposition: {cd}")

        # Inspect DOCX content via python-docx
        docx_bytes = resp.data
        if len(docx_bytes) < 1000:
            _fail("DOCX file size", f"only {len(docx_bytes)} bytes — likely empty")
            return
        _ok("DOCX file size", f"{len(docx_bytes):,} bytes")

        try:
            from docx import Document
            doc = Document(io.BytesIO(docx_bytes))
            all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

            # No English text (allow acronyms like "בע\"מ")
            import re
            long_english = re.findall(r'\b[A-Za-z]{4,}\b', all_text)
            # Filter allowed technical terms
            allowed = {"docx", "DOCX", "True", "False", "None"}
            bad_english = [w for w in long_english if w not in allowed]
            if not bad_english:
                _ok("No unexpected English text in DOCX")
            else:
                _fail("English text found in DOCX", f"e.g.: {bad_english[:5]}")

            # No JSON artifacts
            json_arts = [w for w in ['"claims"', '"amount":', '{"', '}"'] if w in all_text]
            if not json_arts:
                _ok("No JSON artifacts in DOCX")
            else:
                _fail("JSON artifacts in DOCX", str(json_arts))

            # Plaintiff name present
            if "ישראל ישראלי" in all_text or "ישראלי" in all_text:
                _ok("Plaintiff name in DOCX")
            else:
                _fail("Plaintiff name missing from DOCX")

            # Defendant name
            if "טכנולוגיה" in all_text:
                _ok("Defendant name in DOCX")
            else:
                _fail("Defendant name missing from DOCX")

            # Gender correct: "התובע" not "התובעת" and not "התובע/ת"
            if "התובע" in all_text and "התובעת" not in all_text:
                _ok("Gender correct: התובע (male)")
            elif "התובע/ת" in all_text:
                _fail("Gender ambiguity: התובע/ת found — should be resolved to התובע")
            elif "התובעת" in all_text:
                _fail("Wrong gender: התובעת (female) found for male plaintiff")
            else:
                _fail("Gender pronoun 'התובע' not found in DOCX")

            # Check paragraph numbering via DOCX XML (auto-numbering via w:numPr)
            # The generator uses _add_numbering() which sets w:numPr in XML, not "1. " text
            doc_xml = doc.element.xml
            num_pr_count = doc_xml.count("numPr")
            if num_pr_count >= 3:
                _ok("Paragraphs are numbered (XML auto-numbering)",
                    f"{num_pr_count} w:numPr elements found")
            else:
                _fail("Paragraph numbering",
                      f"only {num_pr_count} w:numPr XML elements found")

            # Check calculations appear (formula-like text with ₪)
            shekel_count = all_text.count('₪')
            if shekel_count >= 5:
                _ok("Calculation amounts with ₪ present", f"{shekel_count} occurrences")
            else:
                _fail("Calculations with ₪", f"only {shekel_count} ₪ symbols")

            # Salary amount (18,000) appears
            if "18,000" in all_text or "18000" in all_text:
                _ok("Salary amount 18,000 appears in DOCX")
            else:
                _fail("Salary amount 18,000 missing from DOCX")

            # Check font (David) — inspect XML
            doc_xml = doc.element.xml if hasattr(doc.element, 'xml') else ""
            # RTL direction check in XML
            if "rtl" in doc_xml.lower() or "bidi" in doc_xml.lower():
                _ok("RTL/bidi attributes present in DOCX XML")
            else:
                _fail("RTL attributes missing from DOCX XML")

            _ok("DOCX parsed successfully by python-docx")

        except Exception as e:
            _fail("DOCX parse with python-docx", str(e))
            traceback.print_exc()


# ─────────────────────────────────────────────────────────────────────────────
# TEST 3 — CALCULATION VERIFICATION
# ─────────────────────────────────────────────────────────────────────────────
def test3_calculations():
    section("TEST 3 — Calculation Verification")
    with app.test_client() as c:
        c.post("/login", data={"password": PASSWORD})
        resp = c.post("/calculate",
                      data=json.dumps(TEST_DATA),
                      content_type="application/json")
        result = resp.get_json()
        if not result or not result.get("success"):
            _fail("calculate endpoint", f"error: {result}")
            return
        _ok("calculate endpoint succeeded")

        calcs = result["calculations"]
        dur   = calcs["duration"]
        claims= calcs["claims"]

        # ── Duration
        if dur["total_months"] == EXPECTED_MONTHS:
            _ok("Duration: total_months", f"{dur['total_months']} months")
        else:
            _fail("Duration: total_months",
                  f"expected {EXPECTED_MONTHS}, got {dur['total_months']}")

        if dur["decimal_years"] == EXPECTED_DEC_YEARS:
            _ok("Duration: decimal_years", f"{dur['decimal_years']}")
        else:
            _fail("Duration: decimal_years",
                  f"expected {EXPECTED_DEC_YEARS}, got {dur['decimal_years']}")

        # ── Determining salary
        if calcs["determining_salary"] == EXPECTED_SALARY:
            _ok("Determining salary", f"{calcs['determining_salary']:,.0f} ₪")
        else:
            _fail("Determining salary",
                  f"expected {EXPECTED_SALARY}, got {calcs['determining_salary']}")

        # ── Hourly rate
        if abs(calcs["hourly_rate"] - EXPECTED_HOURLY) < 0.50:
            _ok("Hourly rate", f"{calcs['hourly_rate']:.2f} ₪/h")
        else:
            _fail("Hourly rate",
                  f"expected ≈{EXPECTED_HOURLY:.2f}, got {calcs['hourly_rate']:.2f}")

        # ── Daily rate
        if abs(calcs["daily_rate"] - EXPECTED_DAILY) < 1.0:
            _ok("Daily rate", f"{calcs['daily_rate']:.2f} ₪/day")
        else:
            _fail("Daily rate",
                  f"expected ≈{EXPECTED_DAILY:.2f}, got {calcs['daily_rate']:.2f}")

        # ── Severance (with Section 14)
        sev = claims.get("severance", {})
        if not sev:
            _fail("Severance claim present", "severance key missing")
        else:
            s14 = sev.get("section14")
            if s14:
                _ok("Section 14 data computed")
                if abs(s14.get("statutory", 0) - EXPECTED_SEVERANCE_STATUTORY) < 1:
                    _ok("Severance statutory amount",
                        f"{s14['statutory']:,.0f} ₪  (expected {EXPECTED_SEVERANCE_STATUTORY:,.0f})")
                else:
                    _fail("Severance statutory",
                          f"expected {EXPECTED_SEVERANCE_STATUTORY:,.0f}, got {s14.get('statutory', 0):,.0f}")

                contrib = s14.get("contributions", 0)
                if abs(contrib - EXPECTED_S14_CONTRIB) < 1:
                    _ok("Section 14 contributions (6%)",
                        f"{contrib:,.0f} ₪  (expected {EXPECTED_S14_CONTRIB:,.0f})")
                else:
                    _fail("Section 14 contributions",
                          f"expected {EXPECTED_S14_CONTRIB:,.0f}, got {contrib:,.0f}")

                net = sev.get("amount", 0)
                if abs(net - EXPECTED_SEVERANCE_NET) < 2:
                    _ok("Severance net amount after Section 14",
                        f"{net:,.0f} ₪  (expected {EXPECTED_SEVERANCE_NET:,.0f})")
                else:
                    _fail("Severance net after Section 14",
                          f"expected {EXPECTED_SEVERANCE_NET:,.0f}, got {net:,.0f}")
            else:
                _fail("Section 14 calculation", "section14 key not found in severance claim")

        # ── Prior Notice
        pn = claims.get("prior_notice", {})
        if not pn:
            _fail("Prior notice claim", "prior_notice key missing")
        else:
            if abs(pn.get("amount", 0) - EXPECTED_NOTICE) < 1:
                _ok("Prior notice amount", f"{pn['amount']:,.0f} ₪  (expected {EXPECTED_NOTICE:,.0f})")
            else:
                _fail("Prior notice amount",
                      f"expected {EXPECTED_NOTICE:,.0f}, got {pn.get('amount', 0):,.0f}")
            if pn.get("days") == 30:
                _ok("Prior notice days = 30 (>12 months seniority)")
            else:
                _fail("Prior notice days", f"expected 30, got {pn.get('days')}")

        # ── Vacation
        vac = claims.get("vacation", {})
        if not vac:
            _fail("Vacation claim", "vacation key missing")
        else:
            ent = vac.get("entitled_days", 0)
            if abs(ent - EXPECTED_VAC_TOTAL_DAYS) < 1.0:
                _ok("Vacation entitled days", f"{ent:.2f} days (expected ≈{EXPECTED_VAC_TOTAL_DAYS})")
            else:
                _fail("Vacation entitled days",
                      f"expected ≈{EXPECTED_VAC_TOTAL_DAYS}, got {ent}")

            vac_val = vac.get("amount", 0)
            expected_vac_val = round(EXPECTED_VAC_TOTAL_DAYS * EXPECTED_DAILY, 2)
            if abs(vac_val - expected_vac_val) < 50:
                _ok("Vacation amount",
                    f"{vac_val:,.0f} ₪  (expected ≈{expected_vac_val:,.0f})")
            else:
                _fail("Vacation amount",
                      f"expected ≈{expected_vac_val:,.0f}, got {vac_val:,.0f}")

        # ── Recuperation
        rec = claims.get("recuperation", {})
        if not rec:
            _fail("Recuperation claim", "recuperation key missing")
        else:
            ent_d = rec.get("entitled_days", 0)
            if abs(ent_d - EXPECTED_REC_DAYS) < 0.5:
                _ok("Recuperation days", f"{ent_d:.2f} (expected ≈{EXPECTED_REC_DAYS})")
            else:
                _fail("Recuperation days", f"expected ≈{EXPECTED_REC_DAYS}, got {ent_d}")

            rec_val = rec.get("amount", 0)
            if abs(rec_val - EXPECTED_REC_VALUE) < 50:
                _ok("Recuperation amount",
                    f"{rec_val:,.0f} ₪  (expected ≈{EXPECTED_REC_VALUE:,.0f})")
            else:
                _fail("Recuperation amount",
                      f"expected ≈{EXPECTED_REC_VALUE:,.0f}, got {rec_val:,.0f}")

        # ── Pension gap
        pen = claims.get("pension", {})
        if not pen:
            _fail("Pension claim", "pension key missing")
        else:
            gap = pen.get("amount", 0)
            if abs(gap - EXPECTED_PENSION_GAP) < 2:
                _ok("Pension gap (6 months missing)",
                    f"{gap:,.0f} ₪  (expected {EXPECTED_PENSION_GAP:,.0f})")
            else:
                _fail("Pension gap",
                      f"expected {EXPECTED_PENSION_GAP:,.0f}, got {gap:,.0f}")

        # ── Overtime
        ot = claims.get("overtime", {})
        if not ot:
            _fail("Overtime claim", "overtime key missing")
        else:
            ot_total = ot.get("amount", 0)
            if abs(ot_total - EXPECTED_OT_TOTAL) < 100:
                _ok("Overtime total",
                    f"{ot_total:,.0f} ₪  (expected ≈{EXPECTED_OT_TOTAL:,.0f})")
            else:
                _fail("Overtime total",
                      f"expected ≈{EXPECTED_OT_TOTAL:,.0f}, got {ot_total:,.0f}")

        # ── Salary delay
        sd = claims.get("salary_delay", {})
        if not sd:
            _fail("Salary delay claim", "salary_delay key missing")
        else:
            if abs(sd.get("amount", 0) - EXPECTED_SALARY_DELAY) < 1:
                _ok("Salary delay amount", f"{sd['amount']:,.0f} ₪")
            else:
                _fail("Salary delay amount",
                      f"expected {EXPECTED_SALARY_DELAY:,.0f}, got {sd.get('amount', 0):,.0f}")

        # ── Total
        total = calcs["total"]
        expected_total = sum([
            EXPECTED_SEVERANCE_NET,
            EXPECTED_NOTICE,
            round(EXPECTED_VAC_TOTAL_DAYS * EXPECTED_DAILY, 2),
            EXPECTED_REC_VALUE,
            EXPECTED_PENSION_GAP,
            EXPECTED_OT_TOTAL,
            EXPECTED_SALARY_DELAY,
        ])
        if abs(total - expected_total) < 200:
            _ok("Grand total (approx)", f"{total:,.0f} ₪  (expected ≈{expected_total:,.0f})")
        else:
            _fail("Grand total",
                  f"expected ≈{expected_total:,.0f}, got {total:,.0f}")

        print(f"\n    [CALCULATION SUMMARY]")
        for key, claim in claims.items():
            print(f"    {claim['name']:40s} {claim['amount']:>12,.0f} ₪")
        print(f"    {'סה\"כ':40s} {total:>12,.0f} ₪")

        # Store for later tests
        return calcs


# ─────────────────────────────────────────────────────────────────────────────
# TEST 4 — STATUTE OF LIMITATIONS
# ─────────────────────────────────────────────────────────────────────────────
def test4_limitations():
    section("TEST 4 — Statute of Limitations")
    end_date = date(2026, 1, 15)
    today    = date.today()

    checks = [
        ("severance",    "7 years",  date(2033, 1, 15), "green"),
        ("vacation",     "3 years",  date(2029, 1, 15), "green"),
        ("salary_delay", "60 days",  end_date + timedelta(days=60), None),  # dynamic
    ]

    for key, period, expected_deadline, expected_status in checks:
        lim = _compute_limitation(key, end_date)
        if lim is None:
            _fail(f"Limitation: {key}", "returned None")
            continue

        # Deadline check
        actual_deadline = date.fromisoformat(lim["deadline"])
        if actual_deadline == expected_deadline:
            _ok(f"Limitation deadline: {key}",
                f"{actual_deadline}  ({period})")
        else:
            _fail(f"Limitation deadline: {key}",
                  f"expected {expected_deadline}, got {actual_deadline}")

        # Status check
        actual_status = lim["status"]
        if key == "salary_delay":
            # Dynamic: depends on today vs deadline
            salary_delay_deadline = end_date + timedelta(days=60)
            days_left = (salary_delay_deadline - today).days
            if days_left < 0:
                expected_status = "gray"  # already expired
            elif days_left <= 60:
                expected_status = "red"
            elif days_left <= 180:
                expected_status = "yellow"
            else:
                expected_status = "green"
            note = f"deadline={salary_delay_deadline}, today={today}, days_left={days_left}"
        else:
            note = ""

        if actual_status == expected_status:
            _ok(f"Limitation status: {key}", f"{actual_status}  {note}")
        else:
            _fail(f"Limitation status: {key}",
                  f"expected {expected_status}, got {actual_status}  {note}")

    # Verify via /calculate endpoint that limitations are embedded in claims
    with app.test_client() as c:
        c.post("/login", data={"password": PASSWORD})
        resp = c.post("/calculate",
                      data=json.dumps(TEST_DATA),
                      content_type="application/json")
        result = resp.get_json()
        claims = result.get("calculations", {}).get("claims", {})

        for key in ["severance", "vacation", "salary_delay"]:
            claim = claims.get(key, {})
            lim = claim.get("limitation")
            if lim:
                _ok(f"Limitation embedded in {key} claim",
                    f"deadline={lim['deadline']}, status={lim['status']}")
            else:
                _fail(f"Limitation embedded in {key}", "limitation key missing")

    # Print limitation summary
    print(f"\n    [LIMITATIONS SUMMARY — today is {today}]")
    lim_keys = ["severance", "prior_notice", "vacation", "overtime",
                "pension", "recuperation", "salary_delay"]
    for k in lim_keys:
        lim = _compute_limitation(k, end_date)
        if lim:
            status_color = {"green": GREEN, "yellow": YELLOW,
                            "red": RED, "gray": "\033[90m"}.get(lim["status"], "")
            print(f"    {k:20s} → {lim['deadline']}  "
                  f"{status_color}[{lim['status'].upper()}]{RESET}  "
                  f"({lim['days_remaining']} days remaining)")


# ─────────────────────────────────────────────────────────────────────────────
# TEST 5 — EVIDENCE CHECKLIST (from HTML template)
# ─────────────────────────────────────────────────────────────────────────────
def test5_evidence_checklist():
    section("TEST 5 — Evidence Checklist")
    with app.test_client() as c:
        c.post("/login", data={"password": PASSWORD})
        resp = c.get("/")
        html = resp.data.decode("utf-8")

    # Checklist is rendered by JS in index.html from EVIDENCE_DOCS map
    # Verify the JS evidence map exists and covers our claim types
    evidence_markers = [
        "EVIDENCE_ITEMS",        # JS constant (actual name in index.html)
        "תלושי שכר",             # pay slips — most basic evidence
        "חוזה עבודה",            # employment contract
        "evidence-checklist",    # panel ID
    ]
    for marker in evidence_markers:
        if marker in html:
            _ok(f"Evidence checklist: '{marker}' present in index.html")
        else:
            _fail(f"Evidence checklist: '{marker}' MISSING from index.html")

    # Check evidence items for our specific claim types
    claim_evidence_map = {
        "severance":    ["תלושי שכר", "פיצויי"],
        "vacation":     ["חופשה", "תלושי שכר"],
        "recuperation": ["הבראה"],
        "overtime":     ["שעות נוספות", "נוכחות"],
        "pension":      ["פנסיה", "ביטוח מנהלים"],
    }
    for claim_type, expected_items in claim_evidence_map.items():
        found = any(item in html for item in expected_items)
        if found:
            _ok(f"Evidence items for {claim_type}", f"found one of {expected_items}")
        else:
            _fail(f"Evidence items for {claim_type}", f"none of {expected_items} found")


# ─────────────────────────────────────────────────────────────────────────────
# TEST 6 — DEMAND LETTER
# ─────────────────────────────────────────────────────────────────────────────
def test6_demand_letter():
    section("TEST 6 — Demand Letter")
    with app.test_client() as c:
        c.post("/login", data={"password": PASSWORD})
        resp = c.post(
            "/generate-demand-letter",
            data=json.dumps(TEST_DATA),
            content_type="application/json",
        )

        if resp.status_code == 200:
            _ok("generate-demand-letter HTTP 200")
        else:
            _fail("generate-demand-letter HTTP 200",
                  f"got {resp.status_code}: {resp.data[:200]}")
            return

        ct = resp.content_type
        if "wordprocessingml" in ct or "octet-stream" in ct:
            _ok("Demand letter content-type = DOCX", ct)
        else:
            _fail("Demand letter should be DOCX", f"got {ct}")
            return

        letter_bytes = resp.data
        if len(letter_bytes) < 500:
            _fail("Demand letter file size", f"only {len(letter_bytes)} bytes")
            return
        _ok("Demand letter file size", f"{len(letter_bytes):,} bytes")

        try:
            from docx import Document
            doc = Document(io.BytesIO(letter_bytes))
            all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

            # Employer name
            if "טכנולוגיה" in all_text:
                _ok("Demand letter: employer name present")
            else:
                _fail("Demand letter: employer name missing")

            # Plaintiff name
            if "ישראל ישראלי" in all_text or "ישראלי" in all_text:
                _ok("Demand letter: plaintiff name present")
            else:
                _fail("Demand letter: plaintiff name missing")

            # 14-day deadline
            if "14" in all_text:
                _ok("Demand letter: 14-day deadline mentioned")
            else:
                _fail("Demand letter: 14-day deadline not mentioned")

            # Law references
            law_refs = ["חוק פיצויי פיטורים", "חוק הגנת השכר", "חוק חופשה"]
            found_laws = [l for l in law_refs if l in all_text]
            if found_laws:
                _ok("Demand letter: law references present", str(found_laws))
            else:
                _fail("Demand letter: no law references found")

            # Amount (₪ symbol)
            if "₪" in all_text:
                _ok("Demand letter: amounts with ₪ present")
            else:
                _fail("Demand letter: no ₪ amounts")

            # Firm name / letterhead
            if "לוין תלרז" in all_text or "levin" in all_text.lower():
                _ok("Demand letter: firm name present")
            else:
                _fail("Demand letter: firm name missing")

            # Hebrew text check
            heb = sum(1 for ch in all_text if '\u05d0' <= ch <= '\u05ea')
            if heb > 100:
                _ok("Demand letter is in Hebrew", f"{heb} Hebrew chars")
            else:
                _fail("Demand letter Hebrew", f"only {heb} Hebrew chars")

        except Exception as e:
            _fail("Demand letter DOCX parse", str(e))
            traceback.print_exc()


# ─────────────────────────────────────────────────────────────────────────────
# TEST 7 — SETTLEMENT CALCULATOR (JS-based, verified via HTML presence)
# ─────────────────────────────────────────────────────────────────────────────
def test7_settlement():
    section("TEST 7 — Settlement Calculator")
    with app.test_client() as c:
        c.post("/login", data={"password": PASSWORD})
        resp = c.get("/")
        html = resp.data.decode("utf-8")

    # Check settlement panel and JS logic exist
    settlement_markers = [
        'id="settlement-panel"',
        'settlement-tbody',
        'settlement-strength',
        'renderSettlement',          # JS function
    ]
    for marker in settlement_markers:
        if marker in html:
            _ok(f"Settlement UI: '{marker}' present")
        else:
            _fail(f"Settlement UI: '{marker}' MISSING")

    # Verify three-tier calculation constants exist in JS
    tier_markers = [
        "100",   # full 100%
        "0.7",   # expected ~70%
        "0.5",   # minimum ~50%
    ]
    # Also check for tiered logic
    tier_logic_found = "settlementTiers" in html or "0.7" in html or "settlement_tiers" in html
    if tier_logic_found:
        _ok("Settlement: three-tier logic present (full/expected/minimum)")
    else:
        _fail("Settlement: three-tier logic not found in HTML/JS")

    # Verify the settlement table structure
    if '<table' in html and 'settlement' in html.lower():
        _ok("Settlement: table element present")
    else:
        _fail("Settlement: table element missing")

    # Test calculation logic manually
    # Expected settlement at 70% of total
    with app.test_client() as c:
        c.post("/login", data={"password": PASSWORD})
        resp = c.post("/calculate",
                      data=json.dumps(TEST_DATA),
                      content_type="application/json")
        result = resp.get_json()
        total = result.get("calculations", {}).get("total", 0)

    expected_full    = total
    expected_70pct   = round(total * 0.7, 0)
    expected_50pct   = round(total * 0.5, 0)

    if total > 0:
        _ok("Settlement ranges calculable from total",
            f"Full={expected_full:,.0f} | 70%={expected_70pct:,.0f} | 50%={expected_50pct:,.0f} ₪")
    else:
        _fail("Settlement ranges", "total=0, cannot compute")


# ─────────────────────────────────────────────────────────────────────────────
# TEST 8 — CASE DATABASE
# ─────────────────────────────────────────────────────────────────────────────
def test8_database():
    section("TEST 8 — Case Database")

    # First get calculations
    with app.test_client() as c:
        c.post("/login", data={"password": PASSWORD})
        calc_resp = c.post("/calculate",
                           data=json.dumps(TEST_DATA),
                           content_type="application/json")
        calc_result = calc_resp.get_json()
        calcs = calc_result.get("calculations", {})
        total = calcs.get("total", 0)

        # Save case via API
        case_payload = {
            "plaintiff_name": TEST_DATA["plaintiff_name"],
            "defendant_name": TEST_DATA["defendant_name"],
            "start_date":     TEST_DATA["start_date"],
            "end_date":       TEST_DATA["end_date"],
            "total_amount":   total,
            "stage":          "קליטה",
            "calculations":   calcs,
            "form_data":      TEST_DATA,
            "notes":          "בדיקת QA — ישראל ישראלי",
        }
        save_resp = c.post("/api/cases",
                           data=json.dumps(case_payload),
                           content_type="application/json")

        if save_resp.status_code == 201:
            _ok("Case saved via /api/cases", "HTTP 201")
        else:
            _fail("Case save", f"HTTP {save_resp.status_code}: {save_resp.data[:200]}")
            return

        save_result = save_resp.get_json()
        case_id = save_result.get("id")
        if case_id:
            _ok("Case ID returned", f"id={case_id}")
        else:
            _fail("Case ID", "no id in response")
            return

        # Retrieve and verify
        get_resp = c.get(f"/api/cases/{case_id}")
        if get_resp.status_code == 200:
            _ok("Case retrieved from DB", f"HTTP 200")
        else:
            _fail("Case retrieve", f"HTTP {get_resp.status_code}")
            return

        case_data = get_resp.get_json()

        # Verify all required fields
        checks = [
            ("plaintiff_name", TEST_DATA["plaintiff_name"]),
            ("defendant_name", TEST_DATA["defendant_name"]),
            ("start_date",     TEST_DATA["start_date"]),
            ("end_date",       TEST_DATA["end_date"]),
            ("stage",          "קליטה"),
            ("status",         "active"),
        ]
        for field, expected in checks:
            actual = case_data.get(field)
            if actual == expected:
                _ok(f"DB field: {field}", f"= '{actual}'")
            else:
                _fail(f"DB field: {field}", f"expected '{expected}', got '{actual}'")

        # Total amount
        stored_total = case_data.get("total_amount", 0)
        if abs(stored_total - total) < 1:
            _ok("DB field: total_amount", f"{stored_total:,.0f} ₪")
        else:
            _fail("DB field: total_amount",
                  f"expected {total:,.0f}, got {stored_total:,.0f}")

        # form_data and calculations stored as JSON strings
        fd = case_data.get("form_data")
        if fd:
            try:
                fd_parsed = json.loads(fd) if isinstance(fd, str) else fd
                if fd_parsed.get("plaintiff_name") == TEST_DATA["plaintiff_name"]:
                    _ok("DB field: form_data (JSON round-trip)")
                else:
                    _fail("DB field: form_data", "plaintiff_name mismatch in stored form_data")
            except Exception as e:
                _fail("DB field: form_data JSON parse", str(e))
        else:
            _fail("DB field: form_data", "empty")

        calc_stored = case_data.get("calculations")
        if calc_stored:
            _ok("DB field: calculations stored")
        else:
            _fail("DB field: calculations empty")

        # claim_types
        ct = case_data.get("claim_types", "[]")
        ct_parsed = json.loads(ct) if isinstance(ct, str) else ct
        if isinstance(ct_parsed, list) and len(ct_parsed) >= 5:
            _ok("DB field: claim_types", f"{len(ct_parsed)} claims: {ct_parsed}")
        else:
            _fail("DB field: claim_types", f"got {ct_parsed}")

        return case_id


# ─────────────────────────────────────────────────────────────────────────────
# TEST 9 — KANBAN BOARD / DASHBOARD
# ─────────────────────────────────────────────────────────────────────────────
def test9_kanban(case_id=None):
    section("TEST 9 — Kanban Board / Dashboard")
    with app.test_client() as c:
        c.post("/login", data={"password": PASSWORD})

        # Dashboard page renders
        resp = c.get("/dashboard")
        if resp.status_code == 200:
            _ok("Dashboard page HTTP 200")
        else:
            _fail("Dashboard page", f"HTTP {resp.status_code}")
            return

        html = resp.data.decode("utf-8")

        # Kanban stage columns present
        for stage in ["קליטה", "איסוף מסמכים", "ניתוח ובדיקה", "הוגש", "סגור"]:
            if stage in html:
                _ok(f"Kanban column: '{stage}' present")
            else:
                _fail(f"Kanban column: '{stage}' MISSING")

        # API endpoint returns cases
        api_resp = c.get("/api/cases")
        if api_resp.status_code == 200:
            _ok("GET /api/cases HTTP 200")
        else:
            _fail("GET /api/cases", f"HTTP {api_resp.status_code}")
            return

        cases = api_resp.get_json()
        if not isinstance(cases, list):
            _fail("GET /api/cases returns list", f"got {type(cases)}")
            return
        _ok("GET /api/cases returns list", f"{len(cases)} cases")

        # Find our test case
        our_case = next(
            (c for c in cases if c.get("plaintiff_name") == "ישראל ישראלי"),
            None,
        )
        if our_case:
            _ok("Test case found in dashboard cases list",
                f"id={our_case.get('id')}, stage={our_case.get('stage')}")
        else:
            _fail("Test case in dashboard", "ישראל ישראלי not found in /api/cases")
            return

        # Verify it's in קליטה stage
        if our_case.get("stage") == "קליטה":
            _ok("Test case in correct Kanban column: קליטה")
        else:
            _fail("Test case Kanban stage",
                  f"expected קליטה, got {our_case.get('stage')}")

        # Verify Kanban stage update works
        if case_id:
            patch_resp = c.patch(
                f"/api/cases/{case_id}",
                data=json.dumps({"stage": "ניתוח ובדיקה"}),
                content_type="application/json",
            )
            if patch_resp.status_code == 200:
                _ok("Kanban stage update via PATCH", "קליטה → ניתוח ובדיקה")
                # Verify
                get_resp = c.get(f"/api/cases/{case_id}")
                updated = get_resp.get_json()
                if updated.get("stage") == "ניתוח ובדיקה":
                    _ok("Stage change persisted in DB")
                else:
                    _fail("Stage change", f"got {updated.get('stage')}")
            else:
                _fail("Kanban stage PATCH",
                      f"HTTP {patch_resp.status_code}: {patch_resp.data[:100]}")


# ─────────────────────────────────────────────────────────────────────────────
# TEST 10 — CLIENT PORTAL
# ─────────────────────────────────────────────────────────────────────────────
def test10_client_portal(case_id=None):
    section("TEST 10 — Client Portal")
    if not case_id:
        _fail("Client portal", "No case_id provided — skipping")
        return

    with app.test_client() as c:
        c.post("/login", data={"password": PASSWORD})

        # Generate client token
        token_resp = c.post(f"/api/cases/{case_id}/client-token")
        if token_resp.status_code == 200:
            _ok("Client token generated", f"HTTP 200")
        else:
            _fail("Client token generation",
                  f"HTTP {token_resp.status_code}: {token_resp.data[:100]}")
            return

        token_data = token_resp.get_json()
        token = token_data.get("token")
        if token and len(token) > 20:
            _ok("Client token is valid", f"length={len(token)}")
        else:
            _fail("Client token", f"invalid token: {token}")
            return

        portal_url = token_data.get("url", f"/client/{token}")
        _ok("Client portal URL", portal_url)

        # Access portal WITHOUT login (public endpoint)
        # Create a new client without session
        with app.test_client() as c2:
            portal_resp = c2.get(portal_url)
            if portal_resp.status_code == 200:
                _ok("Client portal accessible without login", "HTTP 200")
            else:
                _fail("Client portal access",
                      f"HTTP {portal_resp.status_code}")
                return

            html = portal_resp.data.decode("utf-8")

            # Should NOT show sensitive legal strategy data
            legal_strategy_markers = [
                '"claims"',     # raw JSON claims data
                'session',      # session data
                '"amount":',    # raw JSON
                "weakness",
                "strategy",
            ]
            for marker in legal_strategy_markers:
                if marker in html:
                    _fail(f"Client portal leaks sensitive data: '{marker}'")
                else:
                    _ok(f"Client portal: '{marker}' not exposed")

            # Should show case status / plaintiff name
            if "ישראל ישראלי" in html:
                _ok("Client portal shows plaintiff name")
            else:
                _fail("Client portal: plaintiff name missing")

            # Should show stage progress
            stage_indicators = ["קליטה", "התקדמות", "שלב", "ניתוח"]
            found_stage = any(s in html for s in stage_indicators)
            if found_stage:
                _ok("Client portal shows case stage/progress")
            else:
                _fail("Client portal: no stage/progress indicator found")

            # Should show next steps
            if "הצעדים הבאים" in html or "הבאים" in html or "המשרד" in html:
                _ok("Client portal shows next steps")
            else:
                _fail("Client portal: no next steps")

            # Evidence checklist in client portal
            if "evidence" in html.lower() or "מסמכים" in html:
                _ok("Client portal shows document checklist")
            else:
                _fail("Client portal: document checklist missing")

        # Verify invalid token returns 404
        with app.test_client() as c3:
            bad_resp = c3.get("/client/invalid-token-xyz-123")
            if bad_resp.status_code == 404:
                _ok("Invalid client token → 404")
            else:
                _fail("Invalid token handling",
                      f"expected 404, got {bad_resp.status_code}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────
def main():
    print(f"\n{BOLD}{'-'*60}")
    print("  LABOR LAW BOT -- FULL QA TEST")
    print(f"  Case: Israel Israeliy vs. Technology Ltd")
    print(f"  Test date: {date.today()}")
    print(f"{'-'*60}{RESET}\n")

    # Ensure test DB is initialized
    _init_db()

    test1_generate_ai()
    test2_docx_generation()
    calcs = test3_calculations()
    test4_limitations()
    test5_evidence_checklist()
    test6_demand_letter()
    test7_settlement()
    case_id = test8_database()
    test9_kanban(case_id)
    test10_client_portal(case_id)

    # ── Final summary ─────────────────────────────────────────────────────
    total_tests = PASS_COUNT + FAIL_COUNT
    print(f"\n{BOLD}{'-'*60}")
    print(f"  QA SUMMARY: {GREEN}{PASS_COUNT} PASSED{RESET}{BOLD}  /  {RED}{FAIL_COUNT} FAILED{RESET}{BOLD}  /  {total_tests} total")
    print(f"{'-'*60}{RESET}")

    if FAIL_COUNT > 0:
        print(f"\n{RED}[FAILURES]{RESET}")
        for r in RESULTS:
            if r[0] == "FAIL":
                print(f"  {RED}✗ {r[1]}{RESET}")
                if len(r) > 2 and r[2]:
                    print(f"      {r[2]}")
    else:
        print(f"\n{GREEN}{BOLD}  All tests passed!{RESET}")

    return FAIL_COUNT


if __name__ == "__main__":
    sys.exit(main())
