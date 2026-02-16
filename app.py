"""
Labor Law Claim Generator Bot - Levin Telraz Law Firm
Generates Israeli labor law claims (כתבי תביעה) based on client intake data.
"""

import json
import math
import logging
from datetime import datetime, date
from dateutil.relativedelta import relativedelta
from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import io

import anthropic

from skill_prompt import SKILL_SYSTEM_PROMPT

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "lt-labor-law-bot-secret-key-2026")
app.config["PERMANENT_SESSION_LIFETIME"] = 86400  # 24 hours in seconds

APP_PASSWORD = os.environ.get("APP_PASSWORD", "LT2026")

# ── Claude API for Legal Text Rewriting ──────────────────────────────────────

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")
_claude_client = None

LEGAL_REWRITE_SYSTEM = (
    "You are an Israeli labor law attorney drafting a כתב תביעה for בית הדין לעבודה. "
    "Rewrite the following facts into professional legal Hebrew suitable for a כתב תביעה. "
    "Use formal legal language, proper clause structure, and reference relevant Israeli labor laws where applicable. "
    "Keep all facts accurate but express them in proper legal drafting style. "
    "Write in third person. Do NOT add any facts that were not provided. "
    "Return ONLY the rewritten legal text, nothing else — no preamble, no explanations."
)


def _get_claude_client():
    """Lazy-init Anthropic client."""
    global _claude_client
    if _claude_client is None and ANTHROPIC_API_KEY:
        _claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    return _claude_client


def rewrite_as_legal_text(raw_text, context=""):
    """Send raw user text to Claude API for professional legal Hebrew rewriting.

    Args:
        raw_text: The user's free-text input to rewrite.
        context: Optional context about the case (names, dates, etc.) to help Claude
                 maintain consistency.

    Returns:
        The rewritten legal text, or the original text if API is unavailable.
    """
    if not raw_text or not raw_text.strip():
        return raw_text

    client = _get_claude_client()
    if client is None:
        logging.warning("ANTHROPIC_API_KEY not set — returning original text without rewriting")
        return raw_text

    user_prompt = raw_text
    if context:
        user_prompt = f"הקשר התיק:\n{context}\n\nהטקסט לשכתוב:\n{raw_text}"

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=2000,
            system=LEGAL_REWRITE_SYSTEM,
            messages=[{"role": "user", "content": user_prompt}],
        )
        return message.content[0].text.strip()
    except Exception as e:
        logging.error(f"Claude API rewrite failed: {e}")
        return raw_text


def generate_full_claim_via_claude(raw_input, structured_data):
    """Send raw facts + structured data to Claude API with SKILL.md system prompt.

    Claude analyzes the input and produces a structured JSON response containing
    the full כתב תביעה: sections, appendices, calculations, and legal citations.

    Args:
        raw_input: The user's raw narrative/facts text.
        structured_data: Dict with all form data (names, dates, salary, gender, claims).

    Returns:
        Parsed JSON dict from Claude, or None if API unavailable/fails.
    """
    client = _get_claude_client()
    if client is None:
        logging.warning("ANTHROPIC_API_KEY not set — AI generation unavailable")
        return None

    # Build the user prompt with all structured data
    gender = structured_data.get("gender", "male")
    gender_label = "זכר" if gender == "male" else "נקבה"

    # Collect selected claims
    claim_keys = {
        "claim_severance": "פיצויי פיטורים",
        "claim_unpaid_salary": "שכר עבודה שלא שולם",
        "claim_overtime": "הפרשי שכר – שעות נוספות",
        "claim_pension": "הפרשי הפרשות לפנסיה",
        "claim_vacation": "הפרשי דמי חופשה ופדיון חופשה",
        "claim_holidays": "דמי חגים והפרשי דמי חג",
        "claim_recuperation": "דמי הבראה",
        "claim_salary_delay": "פיצויי הלנת שכר",
        "claim_emotional": "פיצוי בגין עוגמת נפש",
        "claim_deductions": "ניכויים שלא כדין",
        "claim_documents": "מסירת מסמכי גמר חשבון",
    }
    selected_claims = []
    for key, name in claim_keys.items():
        if structured_data.get(key):
            amount_key = {
                "claim_severance": "severance",
                "claim_unpaid_salary": "unpaid_salary_amount",
                "claim_overtime": "overtime",
                "claim_pension": "pension",
                "claim_vacation": "vacation",
                "claim_holidays": "holidays",
                "claim_recuperation": "recuperation",
                "claim_salary_delay": "salary_delay_amount",
                "claim_emotional": "emotional_amount",
                "claim_deductions": "deduction_amount",
            }.get(key)
            selected_claims.append(name)

    termination_types = {
        "fired": "פיטורים",
        "resigned_justified": "התפטרות בדין מפוטר/ת",
        "resigned": "התפטרות",
    }

    user_prompt = f"""נתוני התיק:
- שם התובע/ת: {structured_data.get('plaintiff_name', '')}
- ת.ז.: {structured_data.get('plaintiff_id', '')}
- מין: {gender_label}
- שם הנתבע/ת: {structured_data.get('defendant_name', '')}
- ח.פ./ע.מ.: {structured_data.get('defendant_id', '')}
- סוג נתבע: {structured_data.get('defendant_type', 'company')}
- בעלים/מנהל: {structured_data.get('defendant_owner', '')}
- תחום עיסוק: {structured_data.get('defendant_business', '')}
- תפקיד: {structured_data.get('job_title', '')}
- תאריך תחילת עבודה: {structured_data.get('start_date', '')}
- תאריך סיום עבודה: {structured_data.get('end_date', '')}
- סוג סיום העסקה: {termination_types.get(structured_data.get('termination_type', 'fired'), 'פיטורים')}
- ימי עבודה בשבוע: {structured_data.get('work_days_per_week', '6')}
- שעות עבודה ביום: {structured_data.get('hours_per_day', '')}
- שכר בסיס: {structured_data.get('base_salary', '')} ₪
- עמלות/תוספות: {structured_data.get('commissions', '0')} ₪
- סדרי עבודה: {structured_data.get('work_schedule', '')}

רכיבי תביעה שנבחרו: {', '.join(selected_claims)}

נתונים נוספים לחישובים:
- צבירת פיצויים בקופה: {structured_data.get('severance_deposited', '0')} ₪
- שכר שלא שולם: {structured_data.get('unpaid_salary_amount', '0')} ₪
- שעות נוספות 125% בשבוע: {structured_data.get('weekly_overtime_125', '0')}
- שעות נוספות 150% בשבוע: {structured_data.get('weekly_overtime_150', '0')}
- שכר שעתי: {structured_data.get('hourly_wage', '0')} ₪
- שעות עבודה סטנדרטיות ביום: {structured_data.get('standard_daily_hours', '8')}
- שעות עבודה בפועל ביום: {structured_data.get('actual_daily_hours', '0')}
- שעות נוספות גלובליות: {structured_data.get('global_ot_hours', '0')}
- הפקדות פנסיה שבוצעו: {structured_data.get('pension_deposited', '0')} ₪
- ימי חופשה ששולמו: {structured_data.get('vacation_days_paid', '0')}
- ימי חג ששולמו: {structured_data.get('holiday_days_paid', '0')}
- ימי הבראה ששולמו: {structured_data.get('recuperation_days_paid', '0')}
- סכום הלנת שכר: {structured_data.get('salary_delay_amount', '0')} ₪
- סכום עוגמת נפש: {structured_data.get('emotional_amount', '25000')} ₪
- סכום ניכויים: {structured_data.get('deduction_amount', '0')} ₪

עובדות גולמיות ותיאור הנסיבות:
{raw_input}
"""

    try:
        message = client.messages.create(
            model="claude-sonnet-4-5-20250929",
            max_tokens=8000,
            system=SKILL_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_prompt}],
        )
        response_text = message.content[0].text.strip()

        # Parse JSON from response (handle possible markdown code fences)
        if response_text.startswith("```"):
            # Strip markdown code fences
            lines = response_text.split("\n")
            # Remove first line (```json) and last line (```)
            json_lines = []
            in_fence = False
            for line in lines:
                if line.strip().startswith("```") and not in_fence:
                    in_fence = True
                    continue
                elif line.strip() == "```" and in_fence:
                    break
                elif in_fence:
                    json_lines.append(line)
            response_text = "\n".join(json_lines)

        return json.loads(response_text)
    except json.JSONDecodeError as e:
        logging.error(f"Claude API returned invalid JSON: {e}")
        logging.debug(f"Raw response: {response_text[:500]}")
        return None
    except Exception as e:
        logging.error(f"Claude API full generation failed: {e}")
        return None


# ── Israeli Labor Law Constants ──────────────────────────────────────────────

MINIMUM_WAGE_2024 = 5880.02  # NIS monthly
MINIMUM_WAGE_HOURLY_2024 = 32.30
PENSION_EMPLOYER_RATE = 0.065  # 6.5%
PENSION_EMPLOYEE_RATE = 0.06   # 6%
SEVERANCE_RATE = 0.0833        # 8.33%
RECUPERATION_DAY_VALUE = 418   # NIS per day (2024)

# Recuperation days entitlement by year
RECUPERATION_DAYS = {
    1: 5,   # Year 1
    2: 6,   # Year 2
    3: 6,   # Year 3
    4: 7,   # Years 4-10
    5: 7,
    6: 7,
    7: 7,
    8: 7,
    9: 7,
    10: 7,
    11: 8,  # Years 11-15
    15: 8,
    16: 9,  # Years 16-19
    20: 10, # Year 20+
}

# Vacation days entitlement by year (6-day work week)
VACATION_DAYS_6DAY = {
    1: 14, 2: 14, 3: 14, 4: 14,
    5: 16,
    6: 18, 7: 21,
    8: 22, 9: 23, 10: 24, 11: 25, 12: 26, 13: 27, 14: 28,
}

# Vacation days entitlement by year (5-day work week)
VACATION_DAYS_5DAY = {
    1: 12, 2: 12, 3: 12, 4: 12,
    5: 13,
    6: 15, 7: 18,
    8: 19, 9: 20, 10: 21, 11: 22, 12: 23, 13: 24, 14: 24,
}

# Jewish holidays per year (typically 9 days)
HOLIDAY_DAYS_PER_YEAR = 9

# Overtime rates
OVERTIME_125_RATE = 0.25  # First 2 hours
OVERTIME_150_RATE = 0.50  # Beyond 2 hours


def calculate_employment_duration(start_date, end_date):
    """Calculate employment duration in years and months."""
    delta = relativedelta(end_date, start_date)
    total_months = delta.years * 12 + delta.months
    years = total_months / 12
    return {
        "years": delta.years,
        "months": delta.months,
        "total_months": total_months,
        "decimal_years": round(years, 2),
    }


def calculate_determining_salary(base_salary, commissions=0, extras=0):
    """Calculate the determining salary (שכר קובע) for labor law purposes."""
    return base_salary + commissions + extras


def calculate_severance(determining_salary, years_decimal):
    """Calculate severance pay (פיצויי פיטורים)."""
    return round(determining_salary * years_decimal, 2)


def calculate_vacation_entitlement(years_decimal, work_days_per_week, daily_rate):
    """Calculate vacation entitlement and monetary value."""
    table = VACATION_DAYS_6DAY if work_days_per_week == 6 else VACATION_DAYS_5DAY
    total_days = 0
    full_years = int(years_decimal)
    fraction = years_decimal - full_years

    for y in range(1, full_years + 1):
        key = min(y, max(table.keys()))
        total_days += table.get(y, table[max(k for k in table if k <= y)])

    if fraction > 0 and full_years + 1 <= max(table.keys()):
        next_year = full_years + 1
        next_entitlement = table.get(next_year, table[max(k for k in table if k <= next_year)])
        total_days += round(next_entitlement * fraction, 2)

    return {
        "total_days": round(total_days, 2),
        "value": round(total_days * daily_rate, 2),
    }


def calculate_recuperation(years_decimal, daily_value=RECUPERATION_DAY_VALUE):
    """Calculate recuperation pay (דמי הבראה)."""
    total_days = 0
    full_years = int(years_decimal)
    fraction = years_decimal - full_years

    for y in range(1, full_years + 1):
        if y <= 3:
            days = RECUPERATION_DAYS.get(y, 6)
        elif y <= 10:
            days = 7
        elif y <= 15:
            days = 8
        elif y <= 19:
            days = 9
        else:
            days = 10
        total_days += days

    if fraction > 0:
        next_y = full_years + 1
        if next_y <= 3:
            days = RECUPERATION_DAYS.get(next_y, 6)
        elif next_y <= 10:
            days = 7
        else:
            days = 8
        total_days += round(days * fraction, 2)

    return {
        "total_days": round(total_days, 2),
        "value": round(total_days * daily_value, 2),
    }


def calculate_holiday_pay(years_decimal, daily_rate, days_paid=0, rate_paid=0):
    """Calculate holiday pay entitlement (דמי חגים)."""
    total_days = round(HOLIDAY_DAYS_PER_YEAR * years_decimal)
    entitled_value = total_days * daily_rate
    paid_value = days_paid * rate_paid
    difference = entitled_value - paid_value
    return {
        "total_days": total_days,
        "entitled_value": round(entitled_value, 2),
        "paid_value": round(paid_value, 2),
        "difference": round(max(0, difference), 2),
    }


def calculate_pension_gaps(monthly_salary, months, employer_rate=PENSION_EMPLOYER_RATE,
                           amount_deposited=0):
    """Calculate pension deposit gaps (הפרשי הפרשות לפנסיה)."""
    total_owed = round(monthly_salary * months * employer_rate, 2)
    gap = round(total_owed - amount_deposited, 2)
    return {
        "total_owed": total_owed,
        "deposited": amount_deposited,
        "gap": max(0, gap),
    }


def calculate_overtime(weekly_overtime_125, weekly_overtime_150, hourly_rate, months):
    """Calculate overtime pay owed (שעות נוספות) - basic weekly input mode."""
    rate_125 = hourly_rate * 1.25
    rate_150 = hourly_rate * 1.50
    surcharge_125 = hourly_rate * OVERTIME_125_RATE
    surcharge_150 = hourly_rate * OVERTIME_150_RATE

    monthly_125 = weekly_overtime_125 * 4 * surcharge_125
    monthly_150 = weekly_overtime_150 * 4 * surcharge_150

    total = round((monthly_125 + monthly_150) * months, 2)
    return {
        "mode": "basic",
        "monthly_125": round(monthly_125, 2),
        "monthly_150": round(monthly_150, 2),
        "total": total,
        "rate_125": round(rate_125, 2),
        "rate_150": round(rate_150, 2),
        "surcharge_125": round(surcharge_125, 2),
        "surcharge_150": round(surcharge_150, 2),
    }


def calculate_overtime_global(hourly_wage, standard_daily_hours, actual_daily_hours,
                              global_ot_hours, work_days_per_week, months):
    """Calculate overtime with global OT comparison (125%/150% daily breakdown)."""
    daily_ot = max(0, actual_daily_hours - standard_daily_hours)
    daily_ot_125 = min(daily_ot, 2)
    daily_ot_150 = max(0, daily_ot - 2)

    rate_125 = hourly_wage * 1.25
    rate_150 = hourly_wage * 1.50

    work_days_per_month = round(work_days_per_week * 4.33, 1)

    monthly_ot_125_hours = round(daily_ot_125 * work_days_per_month, 2)
    monthly_ot_150_hours = round(daily_ot_150 * work_days_per_month, 2)
    monthly_should_pay = round(monthly_ot_125_hours * rate_125 + monthly_ot_150_hours * rate_150, 2)

    global_125_hours = min(global_ot_hours, daily_ot_125 * work_days_per_month) if daily_ot > 0 else min(global_ot_hours, 2 * work_days_per_month)
    global_150_hours = max(0, global_ot_hours - global_125_hours)
    monthly_paid = round(global_125_hours * rate_125 + global_150_hours * rate_150, 2)

    monthly_difference = round(max(0, monthly_should_pay - monthly_paid), 2)
    total = round(monthly_difference * months, 2)

    return {
        "mode": "global",
        "hourly_wage": hourly_wage,
        "standard_daily_hours": standard_daily_hours,
        "actual_daily_hours": actual_daily_hours,
        "daily_ot": round(daily_ot, 2),
        "daily_ot_125": round(daily_ot_125, 2),
        "daily_ot_150": round(daily_ot_150, 2),
        "rate_125": round(rate_125, 2),
        "rate_150": round(rate_150, 2),
        "work_days_per_month": work_days_per_month,
        "monthly_ot_125_hours": monthly_ot_125_hours,
        "monthly_ot_150_hours": monthly_ot_150_hours,
        "monthly_should_pay": monthly_should_pay,
        "global_ot_hours": global_ot_hours,
        "global_125_hours": round(global_125_hours, 2),
        "global_150_hours": round(global_150_hours, 2),
        "monthly_paid": monthly_paid,
        "monthly_difference": monthly_difference,
        "total": total,
        "months": months,
    }


def safe_float(val, default=0):
    """Safely convert a value to float, returning default on failure."""
    if val is None or val == '':
        return default
    try:
        return float(val)
    except (ValueError, TypeError):
        return default


def safe_int(val, default=0):
    """Safely convert a value to int, returning default on failure."""
    if val is None or val == '':
        return default
    try:
        return int(float(val))
    except (ValueError, TypeError):
        return default


def calculate_all_claims(data):
    """Master calculation function for all claim components."""
    start_str = (data.get("start_date") or "").strip()
    end_str = (data.get("end_date") or "").strip()

    if not start_str or not end_str:
        raise ValueError("יש להזין תאריך תחילת עבודה ותאריך סיום עבודה")

    try:
        start = datetime.strptime(start_str, "%Y-%m-%d").date()
    except ValueError:
        raise ValueError(f"תאריך תחילת עבודה אינו תקין: {start_str}")

    try:
        end = datetime.strptime(end_str, "%Y-%m-%d").date()
    except ValueError:
        raise ValueError(f"תאריך סיום עבודה אינו תקין: {end_str}")

    if end <= start:
        raise ValueError("תאריך סיום העבודה חייב להיות מאוחר מתאריך ההתחלה")

    duration = calculate_employment_duration(start, end)

    base_salary = safe_float(data.get("base_salary"), 0)
    commissions = safe_float(data.get("commissions"), 0)
    determining_salary = calculate_determining_salary(base_salary, commissions)

    work_days = safe_int(data.get("work_days_per_week"), 6)
    hours_per_day = safe_float(data.get("hours_per_day"), 8.5 if work_days == 6 else 9)
    monthly_hours = work_days * hours_per_day * 4.33
    hourly_rate = round(determining_salary / monthly_hours, 2) if monthly_hours > 0 else 0
    daily_rate = round(determining_salary / (work_days * 4.33), 2) if work_days > 0 else 0

    results = {
        "duration": duration,
        "determining_salary": determining_salary,
        "hourly_rate": hourly_rate,
        "daily_rate": daily_rate,
        "claims": {},
        "total": 0,
    }

    # Severance (פיצויי פיטורים)
    if data.get("claim_severance"):
        severance = calculate_severance(determining_salary, duration["decimal_years"])
        deposited = safe_float(data.get("severance_deposited"), 0)
        results["claims"]["severance"] = {
            "name": "פיצויי פיטורים",
            "full_amount": severance,
            "deposited": deposited,
            "amount": max(0, round(severance - deposited, 2)),
        }

    # Unpaid salary (שכר שלא שולם)
    if data.get("claim_unpaid_salary"):
        unpaid = safe_float(data.get("unpaid_salary_amount"), 0)
        results["claims"]["unpaid_salary"] = {
            "name": "שכר עבודה שלא שולם",
            "amount": unpaid,
        }

    # Overtime (שעות נוספות)
    if data.get("claim_overtime"):
        ot_hourly_wage = safe_float(data.get("hourly_wage"), 0)
        ot_actual = safe_float(data.get("actual_daily_hours"), 0)
        # Use global OT mode if hourly_wage and actual_daily_hours are filled in
        if ot_hourly_wage > 0 and ot_actual > 0:
            ot_standard = safe_float(data.get("standard_daily_hours"), 8)
            ot_global = safe_float(data.get("global_ot_hours"), 0)
            ot = calculate_overtime_global(
                ot_hourly_wage, ot_standard, ot_actual,
                ot_global, work_days, duration["total_months"],
            )
        else:
            ot = calculate_overtime(
                safe_float(data.get("weekly_overtime_125"), 0),
                safe_float(data.get("weekly_overtime_150"), 0),
                hourly_rate,
                duration["total_months"],
            )
        results["claims"]["overtime"] = {
            "name": "הפרשי שכר – שעות נוספות",
            "amount": ot["total"],
            "details": ot,
        }

    # Pension gaps (הפרשי הפרשות לפנסיה)
    if data.get("claim_pension"):
        pension = calculate_pension_gaps(
            determining_salary,
            duration["total_months"],
            amount_deposited=safe_float(data.get("pension_deposited"), 0),
        )
        results["claims"]["pension"] = {
            "name": "הפרשי הפרשות לפנסיה",
            "amount": pension["gap"],
            "details": pension,
        }

    # Vacation (חופשה)
    if data.get("claim_vacation"):
        vac = calculate_vacation_entitlement(
            duration["decimal_years"], work_days, daily_rate
        )
        paid_days = safe_float(data.get("vacation_days_paid"), 0)
        paid_rate = safe_float(data.get("vacation_rate_paid"), 0)
        paid_value = paid_days * paid_rate
        gap = round(vac["value"] - paid_value, 2)
        results["claims"]["vacation"] = {
            "name": "הפרשי דמי חופשה ופדיון חופשה",
            "entitled_days": vac["total_days"],
            "paid_days": paid_days,
            "amount": max(0, gap),
        }

    # Holiday pay (דמי חגים)
    if data.get("claim_holidays"):
        hol = calculate_holiday_pay(
            duration["decimal_years"],
            daily_rate,
            days_paid=safe_float(data.get("holiday_days_paid"), 0),
            rate_paid=safe_float(data.get("holiday_rate_paid"), 0),
        )
        results["claims"]["holidays"] = {
            "name": "דמי חגים והפרשי דמי חג",
            "amount": hol["difference"],
            "details": hol,
        }

    # Recuperation (הבראה)
    if data.get("claim_recuperation"):
        rec = calculate_recuperation(duration["decimal_years"])
        paid_days = safe_float(data.get("recuperation_days_paid"), 0)
        paid_value = paid_days * RECUPERATION_DAY_VALUE
        gap = round(rec["value"] - paid_value, 2)
        results["claims"]["recuperation"] = {
            "name": "דמי הבראה",
            "entitled_days": rec["total_days"],
            "paid_days": paid_days,
            "amount": max(0, gap),
        }

    # Salary delay damages (פיצויי הלנת שכר)
    if data.get("claim_salary_delay"):
        delay_amount = safe_float(data.get("salary_delay_amount"), 0)
        results["claims"]["salary_delay"] = {
            "name": "פיצויי הלנת שכר",
            "amount": delay_amount,
        }

    # Emotional distress (עוגמת נפש)
    if data.get("claim_emotional"):
        emotional = safe_float(data.get("emotional_amount"), 25000)
        results["claims"]["emotional"] = {
            "name": "פיצוי בגין עוגמת נפש",
            "amount": emotional,
        }

    # Unlawful deductions (ניכויים שלא כדין)
    if data.get("claim_deductions"):
        deductions = safe_float(data.get("deduction_amount"), 0)
        results["claims"]["deductions"] = {
            "name": "ניכויים שלא כדין",
            "amount": deductions,
        }

    # Total
    results["total"] = round(sum(c["amount"] for c in results["claims"].values()), 2)
    return results


def generate_claim_text(data, calculations):
    """Generate the full Hebrew legal claim text based on the firm's template."""

    plaintiff_name = data.get("plaintiff_name", "")
    plaintiff_id = data.get("plaintiff_id", "")
    defendant_name = data.get("defendant_name", "")
    defendant_id = data.get("defendant_id", "")
    defendant_type = data.get("defendant_type", "company")
    defendant_owner = data.get("defendant_owner", "")
    defendant_business = data.get("defendant_business", "")
    job_title = data.get("job_title", "")
    start_date = data.get("start_date", "")
    end_date = data.get("end_date", "")
    termination_type = data.get("termination_type", "fired")
    work_schedule_raw = data.get("work_schedule", "")
    narrative_raw = data.get("narrative", "")

    dur = calculations["duration"]
    det_salary = calculations["determining_salary"]
    hourly = calculations["hourly_rate"]
    daily = calculations["daily_rate"]
    total = calculations["total"]

    # Format dates for Hebrew
    start_dt = datetime.strptime(start_date, "%Y-%m-%d")
    end_dt = datetime.strptime(end_date, "%Y-%m-%d")
    start_fmt = start_dt.strftime("%d.%m.%Y")
    end_fmt = end_dt.strftime("%d.%m.%Y")

    # ── Gender-specific forms ────────────────────────────────────────────
    gender = data.get("gender", "male")
    m = gender == "male"

    # All gendered words used throughout the document
    g = {
        "title": "מר" if m else "הגב'",
        "pronoun": "התובע" if m else "התובעת",
        "he": "הוא" if m else "היא",
        "him": "לו" if m else "לה",
        "his": "שלו" if m else "שלה",
        "worked": "עבד" if m else "עבדה",
        "was_forced": "נאלץ" if m else "נאלצה",
        "represented": "מיוצג" if m else "מיוצגת",
        "submits": "מגיש" if m else "מגישה",
        "worker": "עובד" if m else "עובדת",
        "excellent": "מצוין" if m else "מצוינת",
        "professional": "מקצועי" if m else "מקצועית",
        "performed": "ביצע" if m else "ביצעה",
        "began": "החל" if m else "החלה",
        "hourly_worker": "שעתי" if m else "שעתית",
        "employed": "הועסק" if m else "הועסקה",
        "was": "היה" if m else "היתה",
        "entitled": "זכאי" if m else "זכאית",
        "will_claim": "יטען" if m else "תטען",
        "will_ask": "יבקש" if m else "תבקש",
        "his_salary": "שכרו" if m else "שכרה",
        "his_hourly": "שכרו השעתי" if m else "שכרה השעתי",
        "his_daily": "שכרו היומי" if m else "שכרה היומי",
        "his_monthly": "שכרו החודשי" if m else "שכרה החודשי",
        "his_work": "עבודתו" if m else "עבודתה",
        "his_employment": "העסקתו" if m else "העסקתה",
        "his_rights": "זכויותיו" if m else "זכויותיה",
        "his_seniority": "לוותקו" if m else "לוותקה",
        "fired": "פוטר" if m else "פוטרה",
        "resigned": "התפטר" if m else "התפטרה",
        "his_severance": "פיצויי פיטוריו" if m else "פיצויי פיטוריה",
        "from_him": "ממנו" if m else "ממנה",
        "obligate_him": "לחייבו" if m else "לחייבה",
        "to_hand_him": "למסור לו" if m else "למסור לה",
        "his_possession": "בידי" if m else "בידי",
        "in_his_name": "על שם" if m else "על שם",
        "in_ownership": "שבבעלותו" if m else "שבבעלותה",
        "employer_of": "מעסיקו" if m else "מעסיקה",
        "deducted_from": "משכרו" if m else "משכרה",
        "his_monthly_salary": "משכורתו" if m else "משכורתה",
        "delayed_pay": "שכרו" if m else "שכרה",
        "resigned_as_fired": "להתפטר בדין מפוטר" if m else "להתפטר בדין מפוטרת",
        "was_late": "מאחר" if m else "מאחרת",
        "prevents": "מונע" if m else "מונעת",
    }

    pronoun = g["pronoun"]

    # Defendant label
    if defendant_type == "company":
        defendant_label = f"חברת {defendant_name}"
        defendant_desc = f"הינה חברה בבעלותו ותחת ניהולו של {defendant_owner} העוסקת ב{defendant_business}"
    else:
        defendant_label = defendant_name
        defendant_desc = f"העוסק ב{defendant_business}" if m else f"העוסקת ב{defendant_business}"

    # Termination language
    if termination_type == "fired":
        termination_text = f"עד ש{g['fired']} ביום {end_fmt}"
    elif termination_type == "resigned_justified":
        termination_text = f"עד ש{g['was_forced']} לסיים את {g['his_employment']} בדין {'מפוטר' if m else 'מפוטרת'} ביום {end_fmt}"
    else:
        termination_text = f"עד ש{g['resigned']} ביום {end_fmt}"

    # ── Rewrite free-text fields via Claude API ──────────────────────────
    case_context = (
        f"שם {pronoun}: {plaintiff_name}, ת.ז. {plaintiff_id}\n"
        f"שם הנתבע/ת: {defendant_name}\n"
        f"תקופת העסקה: {start_fmt} עד {end_fmt}\n"
        f"תפקיד: {job_title}\n"
        f"מין: {'זכר' if gender == 'male' else 'נקבה'}\n"
        f"סוג סיום העסקה: {termination_type}"
    )

    work_schedule = work_schedule_raw
    narrative = narrative_raw
    _api_available = _get_claude_client() is not None

    if work_schedule_raw and work_schedule_raw.strip() and _api_available:
        work_schedule = rewrite_as_legal_text(
            f"סדרי העבודה של {pronoun}: {work_schedule_raw}",
            context=case_context,
        )
    if narrative_raw and narrative_raw.strip() and _api_available:
        narrative = rewrite_as_legal_text(narrative_raw, context=case_context)

    sections = []

    # ── Header ──
    sections.append("כ ת ב    ת ב י ע ה")
    sections.append("")

    # ── General ──
    sections.append("כללי")
    sections.append(f"{pronoun} {g['represented']} ע\"י ב\"כ, אשר מענה להמצאת כתבי בית דין הוא, כמצוין בכותרת.")
    sections.append(f"{pronoun} {g['submits']} תביעה זו כנגד הנתבעת בגין הפרת {g['his_rights']} כ{g['worker']} וכאדם, הכול כפי שיפורט להלן.")
    sections.append("הטענות שלהלן הינן חלופיות, מצטברות או משלימות - הכול לפי העניין, הקשר הדברים והדבקם.")
    sections.append("")

    # ── Parties ──
    sections.append("הצדדים")
    sections.append(
        f"{pronoun}, {g['title']} {plaintiff_name}, ת.ז. {plaintiff_id}, "
        f"{g['worked']} בנתבעת החל מיום {start_fmt} {termination_text}, "
        f"סה\"כ {g['worked']} {pronoun} בנתבעת {dur['total_months']} חודשים "
        f"שהם {dur['decimal_years']} שנים (להלן: \"{pronoun}\")."
    )
    sections.append(f"תלושי שכר הנמצאים {g['his_possession']} {pronoun} מצ\"ב ומסומנים כנספח 1.")
    sections.append(
        f"הנתבעת, {defendant_label}, ח.פ./ע.מ. {defendant_id}, "
        f"{defendant_desc} "
        f"ומי ש{g['was']} {g['employer_of']} של {pronoun} בתקופה הרלוונטית לכתב התביעה (להלן: \"הנתבעת\")."
    )
    sections.append("")

    # ── Background ──
    sections.append("רקע עובדתי")
    sections.append(
        f"{pronoun} {g['began']} את {g['his_work']} בנתבעת כ{job_title} החל מיום {start_fmt}."
    )
    if work_schedule:
        if _api_available and work_schedule != work_schedule_raw:
            sections.append(work_schedule)
        else:
            sections.append(f"{g['his_work']} של {pronoun} התנהלה {work_schedule}.")

    sections.append(
        f"לכל אורך תקופת {g['his_employment']}, {pronoun} {g['was']} {g['worker']} {g['excellent']} ו{g['professional']} "
        f"אשר {g['performed']} את {g['his_work']} נאמנה."
    )

    if narrative:
        sections.append("")
        sections.append(narrative)

    sections.append("")

    # ── Employment Scope and Determining Salary ──
    sections.append("היקף משרה ושכר קובע")
    base = safe_float(data.get("base_salary"), 0)
    comm = safe_float(data.get("commissions"), 0)

    salary_desc = f"{g['his_salary']} של {pronoun} עמד על סך של {base:,.0f} ₪ ברוטו"
    if comm > 0:
        salary_desc += f" בגין שכר בסיס ובנוסף {comm:,.0f} ₪ בגין עמלות/תוספות חודשיות"
    salary_desc += "."

    sections.append(salary_desc)
    sections.append(
        f"סה\"כ {g['his_monthly']} הקובע של {pronoun} עמד על {det_salary:,.0f} ₪ ברוטו, "
        f"כך ש{g['his_hourly']} הקובע עמד על סך של {hourly:,.1f} ₪ "
        f"ו{g['his_daily']} הקובע עמד על סך של {daily:,.0f} ₪."
    )
    sections.append("")

    # ── Claim Components ──
    sections.append("רכיבי התביעה")
    sections.append("")

    appendix_num = 2

    claims = calculations["claims"]

    # Unpaid salary
    if "unpaid_salary" in claims:
        c = claims["unpaid_salary"]
        sections.append("שכר עבודה שלא שולם")
        sections.append(
            f"כאמור, {pronoun} {g['will_claim']} כי הנתבעת לא שילמה {g['him']} את {g['his_salary']} כנדרש על פי דין."
        )
        sections.append(
            f"לפיכך, {pronoun} {g['will_ask']} מבית הדין הנכבד לחייב את הנתבעת לשלם ל{pronoun} "
            f"שכר עבודה שלא שולם בסך של {c['amount']:,.0f} ₪ "
            f"בצירוף פיצוי הלנת שכר או הפרשי הצמדה וריבית לפי העניין עד מועד התשלום בפועל."
        )
        sections.append("")

    # Overtime
    if "overtime" in claims:
        c = claims["overtime"]
        d = c["details"]
        sections.append("הפרשי שכר – שעות נוספות")

        if d.get("mode") == "global":
            sections.append(
                f"כאמור, {pronoun} {g['will_claim']} כי הנתבעת לא שילמה {g['him']} כנדרש בגין השעות הנוספות הרבות אותן {g['worked']}."
            )
            sections.append(
                f"{g['his_hourly']} הבסיסי של {pronoun} הינו {d['hourly_wage']:.2f} ₪. "
                f"יום עבודה סטנדרטי: {d['standard_daily_hours']:.1f} שעות. "
                f"שעות עבודה בפועל ביום (ממוצע): {d['actual_daily_hours']:.1f} שעות."
            )
            sections.append(
                f"בהתאם לחוק שעות עבודה ומנוחה, תשי\"א-1951, "
                f"2 השעות הנוספות הראשונות מזכות בתוספת 25% (תעריף {d['rate_125']:.2f} ₪) "
                f"ומעבר לכך בתוספת 50% (תעריף {d['rate_150']:.2f} ₪)."
            )
            sections.append("תחשיב שעות נוספות שהיה צריך לשלם בכל חודש:")
            sections.append(
                f"שעות נוספות ביום: {d['daily_ot']:.1f} שעות "
                f"({d['daily_ot_125']:.1f} שעות × 125% + {d['daily_ot_150']:.1f} שעות × 150%)"
            )
            sections.append(f"ימי עבודה בחודש: {d['work_days_per_month']:.1f} ימים")
            sections.append(
                f"סכום שהיה צריך לשלם בחודש: "
                f"{d['monthly_ot_125_hours']:.1f} שעות × {d['rate_125']:.2f} ₪ + "
                f"{d['monthly_ot_150_hours']:.1f} שעות × {d['rate_150']:.2f} ₪ = "
                f"{d['monthly_should_pay']:,.0f} ₪"
            )
            if d['global_ot_hours'] > 0:
                sections.append(
                    f"שעות נוספות גלובליות ששולמו בפועל: {d['global_ot_hours']:.1f} שעות בחודש "
                    f"(שוויין: {d['monthly_paid']:,.0f} ₪)"
                )
                sections.append(
                    f"הפרש חודשי: {d['monthly_should_pay']:,.0f} ₪ - {d['monthly_paid']:,.0f} ₪ = "
                    f"{d['monthly_difference']:,.0f} ₪"
                )
            sections.append(
                f"הפרש חודשי ({d['monthly_difference']:,.0f} ₪) × {d['months']} חודשי עבודה = "
                f"{c['amount']:,.0f} ₪"
            )
        else:
            sections.append(
                f"כאמור, {pronoun} {g['will_claim']} כי הנתבעת כלל לא שילמה {g['him']} בגין השעות הנוספות הרבות אותן {g['worked']}."
            )
            sections.append(
                f"{g['his_hourly']} של {pronoun} הינו {hourly:.2f} ₪ ומשכך "
                f"תעריף תוספת 25% הינו {d['surcharge_125']:.1f} ₪ "
                f"ותעריף 50% הינו {d['surcharge_150']:.1f} ₪."
            )

        sections.append(
            f"לאור האמור לעיל, בהתאם לתחשיבים, {pronoun} {g['will_ask']} כי בית הדין הנכבד "
            f"יחייב את הנתבעת לשלם ל{pronoun} הפרשי שכר שעות נוספות בסך של {c['amount']:,.0f} ₪ "
            f"בצירוף פיצוי הלנת שכר או הפרשי הצמדה וריבית לפי העניין עד מועד התשלום בפועל."
        )
        sections.append("")

    # Pension
    if "pension" in claims:
        c = claims["pension"]
        sections.append("הפרשי הפרשות לפנסיה")
        sections.append(
            f"בהתאם להוראות צו ההרחבה לפנסיה חובה ולצו ההרחבה בדבר הגדלת ההפרשות לביטוח פנסיוני במשק, "
            f"היה על הנתבעת להפריש ל{pronoun} בגין רכיב תגמולי המעסיק {PENSION_EMPLOYER_RATE*100}% {g['deducted_from']} המלא בכל חודש."
        )
        sections.append(
            f"בהתאם לתחשיבי {pronoun} על הנתבעת לשלם ל{pronoun} הפרשי הפרשות לפנסיה "
            f"בסך {c['amount']:,.0f} ₪."
        )
        sections.append(
            f"לאור האמור לעיל, בהתאם לתחשיבים, {pronoun} {g['will_ask']} כי בית הדין הנכבד "
            f"יחייב את הנתבעת לשלם ל{pronoun} הפרשי הפרשות לפנסיה בסך של {c['amount']:,.0f} ₪ "
            f"בצירוף פיצוי הלנת שכר או הפרשי הצמדה וריבית לפי העניין עד מועד התשלום בפועל."
        )
        sections.append("")

    # Severance
    if "severance" in claims:
        c = claims["severance"]
        sections.append("פיצויי פיטורים")
        if data.get("termination_type") == "resigned_justified":
            sections.append(
                f"{pronoun} {g['will_claim']}, כי לאור ההפרות החמורות והמתמשכות של הנתבעת "
                f"והפגיעה ב{g['his_rights']} הקוגנטיות {g['was_forced']}, בלית ברירה, להודיע על סיום {g['his_employment']}."
            )
            sections.append(
                f"משכך, ובהתאם להוראות חוק פיצויי פיטורים, תשכ\"ג-1963 ולפסיקת בתי הדין לעבודה "
                f"{pronoun} {g['entitled']} {g['resigned_as_fired']} ולמלוא {g['his_severance']}."
            )
        sections.append(
            f"{det_salary:,.0f} ₪ (שכר חודשי קובע) * {dur['decimal_years']} (תקופת העסקה) = {c['full_amount']:,.1f} ₪"
        )
        if c["deposited"] > 0:
            sections.append(f"בניכוי צבירת הפיצויים {g['in_his_name']} {pronoun} בקופה בסך {c['deposited']:,.0f} ₪")
            sections.append(f"סה\"כ {pronoun} {g['entitled']} להשלמת פיצויי פיטורים בסך {c['amount']:,.0f} ₪")
        sections.append(
            f"לאור האמור לעיל, {pronoun} {g['will_ask']} כי בית הדין הנכבד "
            f"יחייב את הנתבעת לשלם ל{pronoun} פיצויי פיטורים בסך של {c['amount']:,.0f} ₪ "
            f"בצירוף פיצוי הלנת שכר או הפרשי הצמדה וריבית לפי העניין עד מועד התשלום בפועל."
        )
        sections.append("")

    # Vacation
    if "vacation" in claims:
        c = claims["vacation"]
        sections.append("הפרשי שכר דמי חופשה ופדיון חופשה")
        sections.append(
            f"בהתאם להוראות חוק חופשה שנתית, תשי\"א-1951 "
            f"{pronoun} {g['was']} {g['entitled']} לצבירת ימי חופשה "
            f"ובהתאם {g['his_seniority']} סה\"כ {c['entitled_days']} ימי חופשה לכל אורך התקופה."
        )
        sections.append(
            f"לאור האמור לעיל {pronoun} {g['will_ask']} כי בית הדין הנכבד "
            f"יחייב את הנתבעת לשלם ל{pronoun} הפרשי שכר דמי חופשה ופדיון חופשה "
            f"בסך של {c['amount']:,.0f} ₪ "
            f"בצירוף פיצוי הלנת שכר או הפרשי הצמדה וריבית לפי העניין עד מועד התשלום בפועל."
        )
        sections.append("")

    # Holidays
    if "holidays" in claims:
        c = claims["holidays"]
        sections.append("דמי חגים והפרשי דמי חג")
        sections.append(
            f"בהתאם להוראות צו ההרחבה הסכם מסגרת 2000 ולאור העובדה כי {pronoun} {g['employed']} "
            f"כ{g['worker']} {g['hourly_worker']}, לאחר 3 חודשי עבודה בנתבעת, {pronoun} {g['was']} {g['entitled']} לתשלום "
            f"בגין {HOLIDAY_DAYS_PER_YEAR} ימי חג בכל שנת עבודה."
        )
        sections.append(
            f"לאור האמור לעיל {pronoun} {g['will_ask']} כי בית הדין הנכבד "
            f"יחייב את הנתבעת לשלם ל{pronoun} דמי חגים והפרשי דמי חג "
            f"בסך של {c['amount']:,.0f} ₪ "
            f"בצירוף פיצוי הלנת שכר או הפרשי הצמדה וריבית לפי העניין עד מועד התשלום בפועל."
        )
        sections.append("")

    # Recuperation
    if "recuperation" in claims:
        c = claims["recuperation"]
        sections.append("דמי הבראה")
        sections.append(
            f"בהתאם להוראות צו ההרחבה בדבר השתתפות המעסיק בהוצאות הבראה ונופש, "
            f"במהלך תקופת {g['his_employment']} {pronoun} {g['was']} {g['entitled']} ל-{c['entitled_days']} ימי הבראה."
        )
        sections.append(
            f"לאור האמור לעיל {pronoun} {g['will_ask']} כי בית הדין הנכבד "
            f"יחייב את הנתבעת לשלם ל{pronoun} דמי הבראה "
            f"בסך של {c['amount']:,.0f} ₪ "
            f"בצירוף פיצוי הלנת שכר או הפרשי הצמדה וריבית לפי העניין עד מועד התשלום בפועל."
        )
        sections.append("")

    # Deductions
    if "deductions" in claims:
        c = claims["deductions"]
        sections.append("ניכויים שלא כדין – תגמולי עובד")
        sections.append(
            f"{pronoun} {g['will_claim']} כי הנתבעת ניכתה {g['deducted_from']} סכומים שלא כדין ובחוסר תום לב."
        )
        sections.append(
            f"לאור האמור לעיל, {pronoun} {g['will_ask']} מבית הדין הנכבד לחייב את הנתבעת "
            f"לשלם ל{pronoun} בגין ניכויים שלא כדין סך של {c['amount']:,.0f} ₪ "
            f"בצירוף פיצוי הלנת שכר או הפרשי הצמדה וריבית לפי העניין עד מועד התשלום בפועל."
        )
        sections.append("")

    # Salary delay
    if "salary_delay" in claims:
        c = claims["salary_delay"]
        sections.append("פיצויי הלנת שכר")
        sections.append(
            f"במרבית תקופת {g['his_employment']} הנתבעת {g['was']} {g['was_late']} באופן שיטתי ועקבי "
            f"בתשלום {g['his_monthly_salary']} החודשית תוך הלנת {g['delayed_pay']} שלא כדין."
        )
        sections.append(
            f"לאור האמור לעיל ובהתאם להוראות חוק הגנת השכר, תשי\"ח-1958 "
            f"הרי ש{pronoun} {g['entitled']} לפיצוי בגין הלנת {g['delayed_pay']} בסך של {c['amount']:,.0f} ₪."
        )
        sections.append("")

    # Emotional distress
    if "emotional" in claims:
        c = claims["emotional"]
        sections.append("פיצוי בגין עוגמת נפש")
        sections.append(
            f"לפיכך, {pronoun} {g['will_ask']} כי בית הדין הנכבד יורה לנתבעת לשלם ל{pronoun} "
            f"פיצוי בגין עוגמת נפש בסך של {c['amount']:,.0f} ₪ "
            f"בצירוף הפרשי הצמדה וריבית ממועד קום העילה ועד לתשלום בפועל."
        )
        sections.append("")

    # ── Document delivery ──
    if data.get("claim_documents"):
        sections.append("מסירת מסמכי גמר חשבון")
        sections.append(
            f"{pronoun} {g['will_claim']} כי חרף העובדה שיחסי העבודה נותקו כבר ביום {end_fmt} "
            f"הנתבעת לא מסרה ל{pronoun} טופס 161 ומסמכי שחרור והעברת בעלות על הקופה {g['in_ownership']} "
            f"ובכך הלכה למעשה {g['prevents']} {g['from_him']} את הגישה לכספי הפנסיה המגיעים {g['him']} על פי דין."
        )
        sections.append(
            f"לאור האמור לעיל {pronoun} {g['will_ask']} כי בית הדין הנכבד יחייב את הנתבעת "
            f"{g['to_hand_him']} את מסמכי גמר החשבון ובהם טופס 161 ערוך על פי דין ומסמכי העברת בעלות."
        )
        sections.append("")

    # ── Summary ──
    sections.append("סיכום")
    sections.append("סיכום רכיבי התביעה:")
    sections.append("")

    for key, claim in claims.items():
        sections.append(f"• {claim['name']}: {claim['amount']:,.0f} ₪")

    sections.append("")
    sections.append(f"סה\"כ סכום התביעה: {total:,.0f} ₪ קרן (לא כולל הצמדה וריבית, שכ\"ט עו\"ד והוצאות)")
    sections.append("")

    sections.append(
        f"לאור ההפרות החמורות של {g['his_rights']} של {pronoun} המתוארות בהרחבה בכתב תביעה זה, "
        f"מתבקש בית הדין הנכבד להזמין את הנתבעת לדין, ו{g['obligate_him']} במלוא סכום התביעה "
        f"בצירוף הפרשי הצמדה וריבית לפי העניין מקום העילה ועד מועד התשלום בפועל "
        f"כמו גם בסעדים ההצהרתיים המבוקשים."
    )
    sections.append(
        f"בנוסף, מתבקש בית הדין הנכבד לחייב את הנתבעת בתשלום הוצאות, שכ\"ט עו\"ד ומע\"מ בגינו."
    )
    sections.append(
        "בית הדין הנכבד מוסמך לדון בתביעה זו לאור מהותה, סכומה, מקום ביצוע העבודה ומענה של הנתבעת."
    )

    return "\n".join(sections)


def generate_claim_text_from_ai(ai_response, data, calculations):
    """Convert Claude's structured AI response into the flat claim text format.

    Takes the AI JSON response and produces a text string compatible with
    generate_docx()'s text parsing logic.

    Args:
        ai_response: Parsed JSON dict from generate_full_claim_via_claude().
        data: Original form data dict.
        calculations: Results from calculate_all_claims().

    Returns:
        A newline-joined string of the claim text.
    """
    sections = []

    sections.append("כ ת ב    ת ב י ע ה")
    sections.append("")

    for section in ai_response.get("sections", []):
        header = section.get("header", "")
        if header:
            sections.append(header)

        for para in section.get("paragraphs", []):
            if para:
                sections.append(para)

        sections.append("")

    # Add summary section using calculated amounts (authoritative source)
    claims = calculations.get("claims", {})
    total = calculations.get("total", 0)

    if claims:
        sections.append("סיכום")
        sections.append("סיכום רכיבי התביעה:")
        sections.append("")

        for key, claim in claims.items():
            sections.append(f"• {claim['name']}: {claim['amount']:,.0f} ₪")

        sections.append("")
        sections.append(
            f"סה\"כ סכום התביעה: {total:,.0f} ₪ קרן (לא כולל הצמדה וריבית, שכ\"ט עו\"ד והוצאות)"
        )
        sections.append("")

    # Final closing paragraphs
    gender = data.get("gender", "male")
    pronoun = "התובע" if gender == "male" else "התובעת"
    g_obligate = "לחייבו" if gender == "male" else "לחייבה"
    g_his_rights = "זכויותיו" if gender == "male" else "זכויותיה"

    sections.append(
        f"לאור ההפרות החמורות של {g_his_rights} של {pronoun} המתוארות בהרחבה בכתב תביעה זה, "
        f"מתבקש בית הדין הנכבד להזמין את הנתבעת לדין, ו{g_obligate} במלוא סכום התביעה "
        f"בצירוף הפרשי הצמדה וריבית לפי העניין מקום העילה ועד מועד התשלום בפועל "
        f"כמו גם בסעדים ההצהרתיים המבוקשים."
    )
    sections.append(
        f"בנוסף, מתבקש בית הדין הנכבד לחייב את הנתבעת בתשלום הוצאות, שכ\"ט עו\"ד ומע\"מ בגינו."
    )
    sections.append(
        "בית הדין הנכבד מוסמך לדון בתביעה זו לאור מהותה, סכומה, מקום ביצוע העבודה ומענה של הנתבעת."
    )

    return "\n".join(sections)


def generate_docx(data, calculations, claim_text):
    """Generate a Word document matching SKILL.md specifications exactly.

    SKILL.md specs:
    - Page: US Letter (12240 × 15840 twips), margins top=709 right=1800 bottom=1276 left=1800
    - Font: David 12pt (24 half-points), RTL bidi, he-IL
    - Numbered paras: ListParagraph, numId, spacing 120/120/360 auto, ind left=-149 right=-709 hanging=425
    - Section headers: bold+underline, NO numbering, ind left=-716 right=-709 firstLine=6
    - Appendix refs: ◄ symbol, bold+underlined, NOT numbered
    - Summary tables: 2-col, bidiVisual BEFORE tblW, last row shaded D9E2F3
    - Signature: 2-col table (spacer 5649 + sig 3377), top border as sig line
    """
    from lxml import etree

    doc = Document()
    WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # ── Page Setup (US Letter, exact twip margins) ───────────────────────
    for section in doc.sections:
        sectPr = section._sectPr
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is None:
            pgSz = etree.SubElement(sectPr, qn('w:pgSz'))
        pgSz.set(qn('w:w'), '12240')
        pgSz.set(qn('w:h'), '15840')

        pgMar = sectPr.find(qn('w:pgMar'))
        if pgMar is None:
            pgMar = etree.SubElement(sectPr, qn('w:pgMar'))
        pgMar.set(qn('w:top'), '709')
        pgMar.set(qn('w:right'), '1800')
        pgMar.set(qn('w:bottom'), '1276')
        pgMar.set(qn('w:left'), '1800')
        pgMar.set(qn('w:header'), '720')
        pgMar.set(qn('w:footer'), '720')

    # ── Configure Default Styles ─────────────────────────────────────────
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'David'
    style_normal.font.size = Pt(12)
    style_normal.font.rtl = True
    # Set cs font on Normal style
    rPr_style = style_normal.element.find(qn('w:rPr'))
    if rPr_style is None:
        rPr_style = etree.SubElement(style_normal.element, qn('w:rPr'))
    rFonts_style = rPr_style.find(qn('w:rFonts'))
    if rFonts_style is None:
        rFonts_style = etree.SubElement(rPr_style, qn('w:rFonts'))
    rFonts_style.set(qn('w:cs'), 'David')
    rFonts_style.set(qn('w:eastAsia'), 'David')
    szCs_style = rPr_style.find(qn('w:szCs'))
    if szCs_style is None:
        szCs_style = etree.SubElement(rPr_style, qn('w:szCs'))
    szCs_style.set(qn('w:val'), '24')

    pf = style_normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style_pPr = style_normal.element.get_or_add_pPr()
    if style_pPr.find(qn('w:bidi')) is None:
        etree.SubElement(style_pPr, qn('w:bidi'))
    # Set spacing on Normal style via XML for exact twip values
    sp = style_pPr.find(qn('w:spacing'))
    if sp is None:
        sp = etree.SubElement(style_pPr, qn('w:spacing'))
    sp.set(qn('w:before'), '120')
    sp.set(qn('w:after'), '120')
    sp.set(qn('w:line'), '360')
    sp.set(qn('w:lineRule'), 'auto')

    # ── Create Numbering ─────────────────────────────────────────────────
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        dummy = doc.add_paragraph('', style='List Number')
        numbering_part = doc.part.numbering_part
        dummy._element.getparent().remove(dummy._element)
    numbering_elm = numbering_part.element

    # SKILL.md numbering: decimal, "%1.", lvlJc="left", b val="0", bCs val="0", lang bidi="he-IL"
    abstract_num_xml = f'''
    <w:abstractNum w:abstractNumId="0" xmlns:w="{WNS}">
        <w:multiLevelType w:val="hybridMultilevel"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="360" w:hanging="360"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="David" w:hAnsi="David" w:cs="David"/>
                <w:b w:val="0"/>
                <w:bCs w:val="0"/>
                <w:sz w:val="24"/>
                <w:szCs w:val="24"/>
                <w:lang w:bidi="he-IL"/>
            </w:rPr>
        </w:lvl>
        <w:lvl w:ilvl="1">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1.%2"/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="720" w:hanging="360"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="David" w:hAnsi="David" w:cs="David"/>
                <w:b w:val="0"/>
                <w:bCs w:val="0"/>
                <w:sz w:val="24"/>
                <w:szCs w:val="24"/>
                <w:lang w:bidi="he-IL"/>
            </w:rPr>
        </w:lvl>
    </w:abstractNum>
    '''
    abstract_num = etree.fromstring(abstract_num_xml)
    numbering_elm.insert(0, abstract_num)

    num_xml = f'''
    <w:num w:numId="2" xmlns:w="{WNS}">
        <w:abstractNumId w:val="0"/>
    </w:num>
    '''
    num_elem = etree.fromstring(num_xml)
    numbering_elm.append(num_elem)

    # ── Helper Functions ─────────────────────────────────────────────────

    def _set_rtl_bidi(p):
        """Set RTL and bidi on a paragraph element."""
        pPr = p._element.get_or_add_pPr()
        if pPr.find(qn('w:bidi')) is None:
            etree.SubElement(pPr, qn('w:bidi'))

    def _set_run_font(run, size=12, bold=False, underline=False, font_name='David'):
        """Configure run font properties including complex script."""
        run.font.name = font_name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.underline = underline
        run.font.rtl = True
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        # bCs for bold complex script
        if bold:
            bCs = rPr.find(qn('w:bCs'))
            if bCs is None:
                etree.SubElement(rPr, qn('w:bCs'))
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = etree.SubElement(rPr, qn('w:szCs'))
        szCs.set(qn('w:val'), str(size * 2))

    def _set_paragraph_spacing(p):
        """Set SKILL.md spacing: before=120, after=120, line=360, lineRule=auto."""
        pPr = p._element.get_or_add_pPr()
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:before'), '120')
        sp.set(qn('w:after'), '120')
        sp.set(qn('w:line'), '360')
        sp.set(qn('w:lineRule'), 'auto')

    def _add_numbering(p, level=0):
        """Add auto-numbering to a paragraph with SKILL.md indentation."""
        pPr = p._element.get_or_add_pPr()
        numPr = etree.SubElement(pPr, qn('w:numPr'))
        ilvl = etree.SubElement(numPr, qn('w:ilvl'))
        ilvl.set(qn('w:val'), str(level))
        numId_el = etree.SubElement(numPr, qn('w:numId'))
        numId_el.set(qn('w:val'), '2')
        # SKILL.md indentation for numbered paras: left=-149, right=-709, hanging=425
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = etree.SubElement(pPr, qn('w:ind'))
        if level == 0:
            ind.set(qn('w:left'), '-149')
            ind.set(qn('w:right'), '-709')
            ind.set(qn('w:hanging'), '425')
        else:
            ind.set(qn('w:left'), '276')
            ind.set(qn('w:right'), '-709')
            ind.set(qn('w:hanging'), '425')

    def add_title(text):
        """Add the main title - centered, bold, large."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_rtl_bidi(p)
        _set_paragraph_spacing(p)
        run = p.add_run(text)
        _set_run_font(run, size=16, bold=True)
        return p

    def add_section_header(text):
        """Add a section header per SKILL.md: bold+underline, NOT numbered, ind left=-716 right=-709 firstLine=6."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_rtl_bidi(p)
        _set_paragraph_spacing(p)
        # SKILL.md section header indentation
        pPr = p._element.get_or_add_pPr()
        ind = etree.SubElement(pPr, qn('w:ind'))
        ind.set(qn('w:left'), '-716')
        ind.set(qn('w:right'), '-709')
        ind.set(qn('w:firstLine'), '6')
        run = p.add_run(text)
        _set_run_font(run, size=12, bold=True, underline=True)
        return p

    def add_numbered_para(text, level=0):
        """Add a numbered body paragraph per SKILL.md."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_rtl_bidi(p)
        _set_paragraph_spacing(p)
        _add_numbering(p, level=level)
        run = p.add_run(text)
        _set_run_font(run, size=12)
        return p

    def add_plain_para(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12,
                       bold=False):
        """Add a plain (non-numbered) paragraph with SKILL.md spacing."""
        p = doc.add_paragraph()
        p.alignment = alignment
        _set_rtl_bidi(p)
        _set_paragraph_spacing(p)
        if text:
            run = p.add_run(text)
            _set_run_font(run, size=size, bold=bold)
        return p

    def add_appendix_ref(text):
        """Add appendix reference per SKILL.md: ◄ symbol, bold+underlined, NOT numbered."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_rtl_bidi(p)
        _set_paragraph_spacing(p)
        pPr = p._element.get_or_add_pPr()
        ind = etree.SubElement(pPr, qn('w:ind'))
        ind.set(qn('w:left'), '-149')
        ind.set(qn('w:right'), '-709')
        # ◄ symbol run (bold, not underlined)
        arrow_run = p.add_run('◄  ')
        _set_run_font(arrow_run, size=12, bold=True, underline=False)
        # Text run (bold + underlined)
        text_run = p.add_run(text)
        _set_run_font(text_run, size=12, bold=True, underline=True)
        return p

    def add_calculation_line(text):
        """Add a calculation/formula line - not numbered."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_rtl_bidi(p)
        _set_paragraph_spacing(p)
        pPr = p._element.get_or_add_pPr()
        ind = etree.SubElement(pPr, qn('w:ind'))
        ind.set(qn('w:left'), '-149')
        ind.set(qn('w:right'), '-709')
        run = p.add_run(text)
        _set_run_font(run, size=12)
        return p

    def set_cell_rtl(cell, text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
        """Set cell text with RTL formatting. No negative indents inside cells."""
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = alignment
        pPr = p._element.get_or_add_pPr()
        if pPr.find(qn('w:bidi')) is None:
            etree.SubElement(pPr, qn('w:bidi'))
        # Ensure no negative indents in cells
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)
        # Compact spacing inside table cells
        sp = pPr.find(qn('w:spacing'))
        if sp is None:
            sp = etree.SubElement(pPr, qn('w:spacing'))
        sp.set(qn('w:before'), '40')
        sp.set(qn('w:after'), '40')
        sp.set(qn('w:line'), '276')
        sp.set(qn('w:lineRule'), 'auto')
        if text:
            for line_idx, line in enumerate(text.split('\n')):
                if line_idx > 0:
                    run = p.add_run()
                    run.add_break()
                run = p.add_run(line)
                _set_run_font(run, size=size, bold=bold)
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            tc.insert(0, tcPr)

    def _make_table_borderless(table):
        """Remove all borders from a table."""
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = etree.SubElement(tbl, qn('w:tblPr'))
        tblBorders = etree.SubElement(tblPr, qn('w:tblBorders'))
        for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = etree.SubElement(tblBorders, qn(f'w:{bn}'))
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
        return tblPr

    def _set_table_bidi(tblPr):
        """Add bidiVisual BEFORE tblW per SKILL.md."""
        # Remove existing bidiVisual if any
        for existing in tblPr.findall(qn('w:bidiVisual')):
            tblPr.remove(existing)
        bidi = etree.SubElement(tblPr, qn('w:bidiVisual'))
        # Move bidiVisual to be before tblW
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is not None:
            tblPr.remove(bidi)
            tblPr.insert(list(tblPr).index(tblW), bidi)
        else:
            # bidiVisual is already at the end, which is fine if no tblW
            pass

    def add_summary_table(claims_dict, total_amount):
        """Add a 2-column summary table with header row, visible borders, blue header."""
        num_rows = len(claims_dict) + 2  # +1 header row, +1 total row
        tbl = doc.add_table(rows=num_rows, cols=2)

        # Set table properties
        tblEl = tbl._element
        tblPr = tblEl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = etree.SubElement(tblEl, qn('w:tblPr'))

        # bidiVisual BEFORE tblW
        etree.SubElement(tblPr, qn('w:bidiVisual'))
        tblW = etree.SubElement(tblPr, qn('w:tblW'))
        tblW.set(qn('w:type'), 'dxa')
        tblW.set(qn('w:w'), '9026')

        # Grid columns: right wider (5513), left narrower (3513)
        tblGrid = tblEl.find(qn('w:tblGrid'))
        if tblGrid is None:
            tblGrid = etree.SubElement(tblEl, qn('w:tblGrid'))
        else:
            for gc in tblGrid.findall(qn('w:gridCol')):
                tblGrid.remove(gc)
        gc1 = etree.SubElement(tblGrid, qn('w:gridCol'))
        gc1.set(qn('w:w'), '5513')
        gc2 = etree.SubElement(tblGrid, qn('w:gridCol'))
        gc2.set(qn('w:w'), '3513')

        # Table borders (all single, visible)
        tblBorders = etree.SubElement(tblPr, qn('w:tblBorders'))
        for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = etree.SubElement(tblBorders, qn(f'w:{bn}'))
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')

        # Row 0: Header row (blue background, white text)
        set_cell_rtl(tbl.rows[0].cells[0], 'רכיב תביעה', bold=True, size=12,
                     alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_rtl(tbl.rows[0].cells[1], 'סכום (₪)', bold=True, size=12,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT)

        # Shade header row blue (1A365D)
        def _shade_cell(cell, fill_color, font_color=None):
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = etree.SubElement(tc, qn('w:tcPr'))
                tc.insert(0, tcPr)
            shd = etree.SubElement(tcPr, qn('w:shd'))
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill_color)
            if font_color:
                for run in cell.paragraphs[0].runs:
                    rPr = run._element.get_or_add_rPr()
                    color = rPr.find(qn('w:color'))
                    if color is None:
                        color = etree.SubElement(rPr, qn('w:color'))
                    color.set(qn('w:val'), font_color)

        _shade_cell(tbl.rows[0].cells[0], '1A365D', 'FFFFFF')
        _shade_cell(tbl.rows[0].cells[1], '1A365D', 'FFFFFF')

        # Data rows: right col = component name (bold, right-aligned), left col = amount
        for i, (key, claim) in enumerate(claims_dict.items()):
            row_idx = i + 1
            set_cell_rtl(tbl.rows[row_idx].cells[0], claim['name'], bold=True, size=12,
                         alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            set_cell_rtl(tbl.rows[row_idx].cells[1], f"{claim['amount']:,.0f} ₪", size=12,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT)

        # Total row (last) - shaded light blue
        last_row = num_rows - 1
        set_cell_rtl(tbl.rows[last_row].cells[0], 'סה"כ', bold=True, size=12,
                     alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_rtl(tbl.rows[last_row].cells[1], f"{total_amount:,.0f} ₪", bold=True, size=12,
                     alignment=WD_ALIGN_PARAGRAPH.LEFT)
        _shade_cell(tbl.rows[last_row].cells[0], 'D9E2F3')
        _shade_cell(tbl.rows[last_row].cells[1], 'D9E2F3')

        return tbl

    # ── Data Extraction ──────────────────────────────────────────────────
    plaintiff_name = data.get("plaintiff_name", "")
    plaintiff_id = data.get("plaintiff_id", "")
    plaintiff_address = data.get("plaintiff_address", "")
    defendant_name = data.get("defendant_name", "")
    defendant_id = data.get("defendant_id", "")
    defendant_address = data.get("defendant_address", "")
    court_name = data.get("court_header", "בית הדין האזורי לעבודה בתל אביב")
    gender = data.get("gender", "male")
    pronoun = "התובע" if gender == "male" else "התובעת"
    total = calculations["total"]
    claims = calculations["claims"]
    attorney_name = data.get("attorney_name", "")
    attorney_id = data.get("attorney_id", "")
    firm_name = data.get("firm_name", "")
    firm_address = data.get("firm_address", "")
    firm_phone = data.get("firm_phone", "")
    firm_fax = data.get("firm_fax", "")
    firm_email = data.get("firm_email", "")

    defendant_label = "הנתבע" if data.get("defendant_type") == "individual" else "הנתבעת"

    # ── Helper: add multiple paragraphs to a cell ────────────────────────
    def set_cell_multiline(cell, lines_spec):
        """Set cell with multiple paragraphs. lines_spec: list of (text, bold, size, alignment) tuples."""
        cell.text = ''
        for idx, (text, bold, size, alignment) in enumerate(lines_spec):
            if idx == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            p.alignment = alignment
            pPr = p._element.get_or_add_pPr()
            if pPr.find(qn('w:bidi')) is None:
                etree.SubElement(pPr, qn('w:bidi'))
            # Remove negative indents
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                pPr.remove(ind)
            # Compact spacing
            sp = pPr.find(qn('w:spacing'))
            if sp is None:
                sp = etree.SubElement(pPr, qn('w:spacing'))
            sp.set(qn('w:before'), '20')
            sp.set(qn('w:after'), '20')
            sp.set(qn('w:line'), '240')
            sp.set(qn('w:lineRule'), 'auto')
            if text:
                run = p.add_run(text)
                _set_run_font(run, size=size, bold=bold)

    # ══════════════════════════════════════════════════════════════════════
    # BUILD THE DOCUMENT — Cover Page (Enbar Shachar format)
    # ══════════════════════════════════════════════════════════════════════

    # Helper: set vertical alignment on a cell
    def _set_cell_valign(cell, val='bottom'):
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            tcPr = etree.SubElement(tc, qn('w:tcPr'))
            tc.insert(0, tcPr)
        va = etree.SubElement(tcPr, qn('w:vAlign'))
        va.set(qn('w:val'), val)

    # ── Table 1: Top Header (INVISIBLE borders, NO bidiVisual) ─────────
    # Without bidiVisual: cell[0]=LEFT, cell[1]=RIGHT
    # LEFT = court name (2 lines), RIGHT = סע"ש / בפני
    hdr_tbl = doc.add_table(rows=1, cols=2)
    hdr_el = hdr_tbl._element
    hdr_tblPr = hdr_el.find(qn('w:tblPr'))
    if hdr_tblPr is None:
        hdr_tblPr = etree.SubElement(hdr_el, qn('w:tblPr'))

    hdr_tblW = etree.SubElement(hdr_tblPr, qn('w:tblW'))
    hdr_tblW.set(qn('w:type'), 'dxa')
    hdr_tblW.set(qn('w:w'), '9026')
    _make_table_borderless(hdr_tbl)

    hdr_grid = hdr_el.find(qn('w:tblGrid'))
    if hdr_grid is None:
        hdr_grid = etree.SubElement(hdr_el, qn('w:tblGrid'))
    else:
        for gc in hdr_grid.findall(qn('w:gridCol')):
            hdr_grid.remove(gc)
    for w in ['4513', '4513']:
        gc = etree.SubElement(hdr_grid, qn('w:gridCol'))
        gc.set(qn('w:w'), w)

    # Parse court name: split "בית הדין האזורי לעבודה בתל אביב" into 2 lines
    court_base = court_name
    court_location = ""
    if " ב" in court_name:
        # Split at last " ב" which starts the location (e.g. "בתל אביב")
        parts = court_name.rsplit(" ב", 1)
        if len(parts) == 2:
            court_base = parts[0]
            court_location = "ב" + parts[1]

    # cell[0] = LEFT side: court name on 2 lines
    court_lines = [(court_base, True, 12, WD_ALIGN_PARAGRAPH.LEFT)]
    if court_location:
        court_lines.append((court_location, True, 12, WD_ALIGN_PARAGRAPH.LEFT))
    set_cell_multiline(hdr_tbl.rows[0].cells[0], court_lines)

    # cell[1] = RIGHT side: סע"ש and בפני
    set_cell_multiline(hdr_tbl.rows[0].cells[1], [
        ('סע"ש ________', False, 11, WD_ALIGN_PARAGRAPH.RIGHT),
        ('בפני _________', False, 11, WD_ALIGN_PARAGRAPH.RIGHT),
    ])

    # ── Table 2: Parties Section (VISIBLE borders) ───────────────────────
    # Rows: בעניין, plaintiff, נגד, defendant, מהות/סכום
    # 2 columns (bidiVisual): col0=content (wide RIGHT), col1=label (narrow LEFT)
    parties_tbl = doc.add_table(rows=5, cols=2)
    pt_el = parties_tbl._element
    pt_tblPr = pt_el.find(qn('w:tblPr'))
    if pt_tblPr is None:
        pt_tblPr = etree.SubElement(pt_el, qn('w:tblPr'))

    etree.SubElement(pt_tblPr, qn('w:bidiVisual'))
    pt_tblW = etree.SubElement(pt_tblPr, qn('w:tblW'))
    pt_tblW.set(qn('w:type'), 'dxa')
    pt_tblW.set(qn('w:w'), '9026')

    pt_borders = etree.SubElement(pt_tblPr, qn('w:tblBorders'))
    for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = etree.SubElement(pt_borders, qn(f'w:{bn}'))
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')

    pt_grid = pt_el.find(qn('w:tblGrid'))
    if pt_grid is None:
        pt_grid = etree.SubElement(pt_el, qn('w:tblGrid'))
    else:
        for gc in pt_grid.findall(qn('w:gridCol')):
            pt_grid.remove(gc)
    for w in ['7026', '2000']:
        gc = etree.SubElement(pt_grid, qn('w:gridCol'))
        gc.set(qn('w:w'), w)

    # Row 0: "בעניין:"
    set_cell_rtl(parties_tbl.rows[0].cells[0], 'בעניין:', bold=True, size=12,
                 alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    set_cell_rtl(parties_tbl.rows[0].cells[1], '', size=11)

    # Row 1: Plaintiff details (col0) | label (col1)
    plaintiff_lines = []
    name_id = f'{plaintiff_name}, ת.ז. {plaintiff_id}' if plaintiff_id else plaintiff_name
    plaintiff_lines.append((name_id, True, 12, WD_ALIGN_PARAGRAPH.RIGHT))
    if plaintiff_address:
        plaintiff_lines.append((plaintiff_address, False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    if attorney_name:
        plaintiff_lines.append((f'באמצעות ב"כ עוה"ד {attorney_name}', False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    if firm_name:
        plaintiff_lines.append((f'ממשרד {firm_name}' if not firm_name.startswith('ממשרד') else firm_name, False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    if firm_address:
        # Split address into street line and building/floor line
        addr_parts = firm_address.split(',')
        for part in addr_parts:
            part = part.strip()
            if part:
                plaintiff_lines.append((part, False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    contact_parts = []
    if firm_phone:
        contact_parts.append(f"טל': {firm_phone}")
    if firm_fax:
        contact_parts.append(f"פקסי': {firm_fax}")
    if contact_parts:
        plaintiff_lines.append((' '.join(contact_parts), False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    if firm_email:
        plaintiff_lines.append((firm_email, False, 11, WD_ALIGN_PARAGRAPH.RIGHT))

    set_cell_multiline(parties_tbl.rows[1].cells[0], plaintiff_lines)
    set_cell_rtl(parties_tbl.rows[1].cells[1], pronoun, bold=True, size=12,
                 alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_valign(parties_tbl.rows[1].cells[1], 'bottom')

    # Row 2: "- נגד -" centered
    set_cell_rtl(parties_tbl.rows[2].cells[0], '- נגד -', bold=True, size=12,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_rtl(parties_tbl.rows[2].cells[1], '', size=11)

    # Row 3: Defendant details (col0) | label (col1)
    defendant_lines = []
    defendant_lines.append((defendant_name, True, 12, WD_ALIGN_PARAGRAPH.RIGHT))
    if defendant_id:
        defendant_lines.append((f'ח.פ {defendant_id}', False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    if defendant_address:
        defendant_lines.append((defendant_address, False, 11, WD_ALIGN_PARAGRAPH.RIGHT))

    set_cell_multiline(parties_tbl.rows[3].cells[0], defendant_lines)
    set_cell_rtl(parties_tbl.rows[3].cells[1], defendant_label, bold=True, size=12,
                 alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_valign(parties_tbl.rows[3].cells[1], 'bottom')

    # Row 4: מהות/סכום inside the bordered parties table
    amount_str = f'{total:,.0f} ₪'
    set_cell_multiline(parties_tbl.rows[4].cells[0], [
        (f'מהות התביעה: הצהרתית וכספית', True, 11, WD_ALIGN_PARAGRAPH.RIGHT),
        (f'סכום התביעה: {amount_str}', True, 11, WD_ALIGN_PARAGRAPH.RIGHT),
    ])
    set_cell_rtl(parties_tbl.rows[4].cells[1], '', size=11)

    # ── Header Summary Table (financial breakdown) ───────────────────────
    add_plain_para('')
    add_summary_table(claims, total)

    # ── Title ────────────────────────────────────────────────────────────
    add_title('כ ת ב    ת ב י ע ה')

    # ── Body - Parse claim_text and format properly ──────────────────────
    section_headers = {
        "כללי", "הצדדים", "רקע עובדתי", "היקף משרה ושכר קובע",
        "רכיבי התביעה", "סיכום",
        "שכר עבודה שלא שולם", "הפרשי שכר – שעות נוספות",
        "הפרשי הפרשות לפנסיה", "פיצויי פיטורים",
        "הפרשי שכר דמי חופשה ופדיון חופשה",
        "דמי חגים והפרשי דמי חג", "דמי הבראה",
        "ניכויים שלא כדין – תגמולי עובד", "פיצויי הלנת שכר",
        "פיצוי בגין עוגמת נפש", "מסירת מסמכי גמר חשבון",
        "עילות התביעה", "הסעדים המבוקשים",
        "תחשיב שעות נוספות שהיה צריך לשלם בכל חודש:",
    }

    lines = claim_text.split("\n")
    in_summary = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        elif stripped == "כ ת ב    ת ב י ע ה":
            continue
        elif stripped == "סיכום רכיבי התביעה:":
            add_section_header(stripped)
            in_summary = True
            continue
        elif in_summary and stripped.startswith("•"):
            continue  # Skip bullet items; we'll use summary table instead
        elif in_summary and not stripped.startswith("•") and "סה\"כ סכום התביעה" not in stripped:
            in_summary = False
            # Fall through to normal processing
        elif "סה\"כ סכום התביעה" in stripped:
            continue  # Skip; shown in summary table
        if stripped in section_headers:
            add_section_header(stripped)
        elif stripped.startswith("תלושי שכר") and "נספח" in stripped:
            add_appendix_ref(stripped)
        elif any(c in stripped for c in ['=', '×']) and '₪' in stripped:
            add_calculation_line(stripped)
        else:
            add_numbered_para(stripped)

    # ── End Summary Table (must match header summary) ────────────────────
    add_section_header("סיכום רכיבי התביעה")
    add_summary_table(claims, total)

    add_plain_para(
        f'סה"כ סכום התביעה: {total:,.0f} ₪ קרן (לא כולל הצמדה וריבית, שכ"ט עו"ד והוצאות)',
        bold=True
    )

    # Final legal paragraphs from claim text (after the summary section)
    found_summary_end = False
    for line in lines:
        stripped = line.strip()
        if "סה\"כ סכום התביעה" in stripped:
            found_summary_end = True
            continue
        if found_summary_end and stripped:
            add_numbered_para(stripped)

    # ── Power of Attorney Note ───────────────────────────────────────────
    add_appendix_ref('ייפוי כוח מצורף לכתב התביעה')

    # ── Signature Table (2-col: spacer 5649 + sig 3377, per SKILL.md) ────
    add_plain_para('')

    sig_table = doc.add_table(rows=1, cols=2)
    sig_tbl_el = sig_table._element
    sig_tblPr = sig_tbl_el.find(qn('w:tblPr'))
    if sig_tblPr is None:
        sig_tblPr = etree.SubElement(sig_tbl_el, qn('w:tblPr'))

    # bidiVisual BEFORE tblW
    sig_bidi = etree.SubElement(sig_tblPr, qn('w:bidiVisual'))
    sig_tblW = etree.SubElement(sig_tblPr, qn('w:tblW'))
    sig_tblW.set(qn('w:type'), 'dxa')
    sig_tblW.set(qn('w:w'), '9026')

    # Remove borders
    sig_borders = etree.SubElement(sig_tblPr, qn('w:tblBorders'))
    for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = etree.SubElement(sig_borders, qn(f'w:{bn}'))
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')

    # Grid: spacer 5649 + sig 3377
    sig_grid = sig_tbl_el.find(qn('w:tblGrid'))
    if sig_grid is None:
        sig_grid = etree.SubElement(sig_tbl_el, qn('w:tblGrid'))
    else:
        for gc in sig_grid.findall(qn('w:gridCol')):
            sig_grid.remove(gc)
    gc1 = etree.SubElement(sig_grid, qn('w:gridCol'))
    gc1.set(qn('w:w'), '5649')
    gc2 = etree.SubElement(sig_grid, qn('w:gridCol'))
    gc2.set(qn('w:w'), '3377')

    # Spacer cell (empty)
    set_cell_rtl(sig_table.rows[0].cells[0], '', size=12)

    # Signature cell with top border (signature line)
    sig_cell = sig_table.rows[0].cells[1]
    if attorney_name and attorney_id:
        sig_text = f'{attorney_name}, עו"ד\nמ.ר. {attorney_id}\nב"כ {pronoun}'
    else:
        sig_text = f'__________________\nב"כ {pronoun}'
    set_cell_rtl(sig_cell, sig_text, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Add top border to signature cell (serves as signature line)
    sig_tc = sig_cell._element
    sig_tcPr = sig_tc.find(qn('w:tcPr'))
    if sig_tcPr is None:
        sig_tcPr = etree.SubElement(sig_tc, qn('w:tcPr'))
        sig_tc.insert(0, sig_tcPr)
    sig_tcBorders = etree.SubElement(sig_tcPr, qn('w:tcBorders'))
    top_border = etree.SubElement(sig_tcBorders, qn('w:top'))
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '4')
    top_border.set(qn('w:space'), '0')
    top_border.set(qn('w:color'), 'auto')

    return doc


# ── Flask Routes ─────────────────────────────────────────────────────────────

@app.before_request
def require_login():
    allowed = ("login", "static", "service_worker", "manifest")
    if request.endpoint not in allowed and not session.get("authenticated"):
        # Return JSON error for AJAX/API requests instead of HTML redirect
        if request.is_json or request.headers.get("Accept", "").startswith("application/json"):
            return jsonify({"success": False, "error": "Session expired — please refresh and log in again"}), 401
        return redirect(url_for("login"))


@app.errorhandler(500)
def internal_error(e):
    """Return JSON for API errors instead of HTML error page."""
    if request.is_json or request.headers.get("Accept", "").startswith("application/json"):
        return jsonify({"success": False, "error": f"Internal server error: {e}"}), 500
    return f"<h1>500 Internal Server Error</h1><p>{e}</p>", 500


@app.errorhandler(404)
def not_found(e):
    """Return JSON for API 404s instead of HTML."""
    if request.is_json or request.headers.get("Accept", "").startswith("application/json"):
        return jsonify({"success": False, "error": "Route not found"}), 404
    return redirect(url_for("login"))


@app.route("/sw.js")
def service_worker():
    return app.send_static_file("sw.js"), 200, {"Content-Type": "application/javascript", "Service-Worker-Allowed": "/"}


@app.route("/manifest.json")
def manifest():
    return app.send_static_file("manifest.json"), 200, {"Content-Type": "application/manifest+json"}


@app.route("/login", methods=["GET", "POST"])
def login():
    error = None
    if request.method == "POST":
        password = request.form.get("password", "")
        if password == APP_PASSWORD:
            session.permanent = True
            session["authenticated"] = True
            return redirect(url_for("index"))
        else:
            error = "סיסמה שגויה, נסה שוב"
    return render_template("login.html", error=error)


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/calculate", methods=["POST"])
def calculate():
    data = request.json
    try:
        calculations = calculate_all_claims(data)
        claim_text = generate_claim_text(data, calculations)
        return jsonify({
            "success": True,
            "calculations": calculations,
            "claim_text": claim_text,
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/generate-ai", methods=["POST"])
def generate_ai_route():
    """AI-powered full claim generation using SKILL.md system prompt."""
    data = request.json
    raw_text = data.get("raw_text", "")

    if not raw_text or not raw_text.strip():
        return jsonify({"success": False, "error": "יש להזין עובדות גולמיות לטקסט"}), 400

    try:
        # Calculate claims using existing logic (for amounts and fallback)
        calculations = calculate_all_claims(data)

        # Generate via Claude AI
        ai_response = generate_full_claim_via_claude(raw_text, data)

        if ai_response is None:
            # Fallback to template-based generation
            claim_text = generate_claim_text(data, calculations)
            return jsonify({
                "success": True,
                "mode": "template",
                "calculations": calculations,
                "claim_text": claim_text,
                "ai_response": None,
            })

        # Generate claim text from AI response for preview
        claim_text = generate_claim_text_from_ai(ai_response, data, calculations)

        return jsonify({
            "success": True,
            "mode": "ai",
            "calculations": calculations,
            "claim_text": claim_text,
            "ai_response": ai_response,
        })
    except Exception as e:
        logging.error(f"AI generation route error: {e}")
        return jsonify({"success": False, "error": str(e)}), 400


@app.route("/generate-docx", methods=["POST"])
def generate_docx_route():
    data = request.json
    try:
        calculations = calculate_all_claims(data)

        # Check if AI response is provided (from /generate-ai flow)
        ai_response = data.get("_ai_response")
        if ai_response:
            claim_text = generate_claim_text_from_ai(ai_response, data, calculations)
        else:
            claim_text = generate_claim_text(data, calculations)

        doc = generate_docx(data, calculations, claim_text)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        plaintiff = data.get("plaintiff_name", "claim")
        filename = f"כתב_תביעה_{plaintiff}.docx"

        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 400


if __name__ == "__main__":
    import os
    debug = os.environ.get("FLASK_DEBUG", "true").lower() == "true"
    port = int(os.environ.get("PORT", 5000))
    app.run(debug=debug, port=port, host="0.0.0.0")
