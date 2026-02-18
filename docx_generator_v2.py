"""
docx_generator_v2.py — Clean implementation of כתב תביעה DOCX generator.

Receives form_data dict + plain Hebrew text from Claude, produces a .docx file.
All formatting specs from SKILL.md are hardcoded as constants below.
"""

import re
import logging
from lxml import etree

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from claude_stages import fix_gender

# ══════════════════════════════════════════════════════════════════════════════
# FORMATTING CONSTANTS (from SKILL.md)
# ══════════════════════════════════════════════════════════════════════════════

FONT_NAME = "David"
FONT_SIZE_PT = 12
FONT_SIZE_HALF_POINTS = 24  # 12pt = 24 half-points

PAGE_WIDTH = 12240   # US Letter width in twips
PAGE_HEIGHT = 15840  # US Letter height in twips

MARGIN_TOP = 709
MARGIN_RIGHT = 1800
MARGIN_BOTTOM = 1276
MARGIN_LEFT = 1800
MARGIN_HEADER = 720
MARGIN_FOOTER = 720

LINE_SPACING = 360       # 1.5 line spacing
LINE_RULE = "auto"
PARA_BEFORE = 120
PARA_AFTER = 120

# Numbered paragraph indents
NUM_INDENT_LEFT = -149
NUM_INDENT_RIGHT = -709
NUM_INDENT_HANGING = 425

# Section header indents
HDR_INDENT_LEFT = -716
HDR_INDENT_RIGHT = -709
HDR_INDENT_FIRSTLINE = 6

# Table widths
TABLE_WIDTH = 9026
SUMMARY_COL_RIGHT = 5513  # component name column
SUMMARY_COL_LEFT = 3513   # amount column
SIG_COL_SPACER = 5649
SIG_COL_SIG = 3377

# Colors
HEADER_BG = "1A365D"
HEADER_FG = "FFFFFF"
TOTAL_ROW_BG = "D9E2F3"

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ══════════════════════════════════════════════════════════════════════════════
# MAIN FUNCTION
# ══════════════════════════════════════════════════════════════════════════════

def generate_claim_docx(form_data: dict, ai_text: str, output_path: str):
    """
    Generate a כתב תביעה .docx from form data and plain AI text.

    form_data: contains plaintiff name, defendant name, dates, salary, gender, amounts, etc.
    ai_text: plain Hebrew text from Claude, sections separated by === TITLE ===
    output_path: where to save the .docx
    """
    doc = Document()

    # Step 0: Setup document formatting
    _setup_page(doc)
    _setup_styles(doc)
    _setup_numbering(doc)

    # Extract common data
    gender = form_data.get("gender", "male")
    pronoun = "התובע" if gender == "male" else "התובעת"
    claims = form_data.get("_claims", {})
    total = form_data.get("_total", 0)

    # Step 1: Cover page
    _build_cover_page(doc, form_data, claims, total)

    # Step 2: Parse ai_text into sections
    sections = _parse_sections(ai_text)
    logging.info(f"docx_generator_v2: parsed {len(sections)} sections from AI text")

    # Step 3: Write sections into document body
    for sec in sections:
        title = sec["title"]
        lines = sec["lines"]

        # Skip סיכום — we render our own summary table
        if title and "סיכום" in title:
            logging.info(f"  Skipping summary section: '{title}'")
            continue

        # Section header (bold+underline, not numbered)
        if title and _has_hebrew(title):
            _add_section_header(doc, title)

        # Content lines
        for line in lines:
            line = _clean_line(line)
            if not line:
                continue
            if not _has_hebrew(line):
                continue

            if line.startswith("◄"):
                _add_appendix_ref(doc, line.lstrip("◄ "))
            elif "₪" in line and any(c in line for c in ["=", "×"]):
                _add_calculation_line(doc, line)
            else:
                _add_numbered_para(doc, line)

    # Step 4: Summary table
    _add_section_header(doc, "סיכום רכיבי התביעה")
    _add_summary_table(doc, claims, total)
    _add_plain_para(doc,
        f'סה"כ סכום התביעה: {total:,.0f} ₪ קרן (לא כולל הצמדה וריבית, שכ"ט עו"ד והוצאות)',
        bold=True)

    # Closing paragraphs
    g_obligate = "לחייבו" if gender == "male" else "לחייבה"
    g_rights = "זכויותיו" if gender == "male" else "זכויותיה"
    _add_numbered_para(doc,
        f"לאור ההפרות החמורות של {g_rights} של {pronoun} המתוארות בהרחבה בכתב תביעה זה, "
        f"מתבקש בית הדין הנכבד להזמין את הנתבעת לדין, ו{g_obligate} במלוא סכום התביעה "
        f"בצירוף הפרשי הצמדה וריבית לפי העניין מקום העילה ועד מועד התשלום בפועל "
        f"כמו גם בסעדים ההצהרתיים המבוקשים.")
    _add_numbered_para(doc,
        f'בנוסף, מתבקש בית הדין הנכבד לחייב את הנתבעת בתשלום הוצאות, שכ"ט עו"ד ומע"מ בגינו.')
    _add_numbered_para(doc,
        "בית הדין הנכבד מוסמך לדון בתביעה זו לאור מהותה, סכומה, מקום ביצוע העבודה ומענה של הנתבעת.")
    _add_appendix_ref(doc, "ייפוי כוח מצורף לכתב התביעה")

    # Step 5: Signature block
    _add_signature_block(doc, form_data)

    # Step 6: Gender replacements
    _apply_gender_to_doc(doc, gender)

    # Step 7: Final Hebrew proofing pass — ensure every run has rtl + lang
    _ensure_hebrew_proofing(doc)

    # Step 8: Save
    doc.save(output_path)
    logging.info(f"docx_generator_v2: saved to {output_path}")


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT SETUP HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _setup_page(doc):
    """Configure page size and margins per SKILL.md."""
    for section in doc.sections:
        sectPr = section._sectPr
        pgSz = sectPr.find(qn("w:pgSz"))
        if pgSz is None:
            pgSz = etree.SubElement(sectPr, qn("w:pgSz"))
        pgSz.set(qn("w:w"), str(PAGE_WIDTH))
        pgSz.set(qn("w:h"), str(PAGE_HEIGHT))

        pgMar = sectPr.find(qn("w:pgMar"))
        if pgMar is None:
            pgMar = etree.SubElement(sectPr, qn("w:pgMar"))
        pgMar.set(qn("w:top"), str(MARGIN_TOP))
        pgMar.set(qn("w:right"), str(MARGIN_RIGHT))
        pgMar.set(qn("w:bottom"), str(MARGIN_BOTTOM))
        pgMar.set(qn("w:left"), str(MARGIN_LEFT))
        pgMar.set(qn("w:header"), str(MARGIN_HEADER))
        pgMar.set(qn("w:footer"), str(MARGIN_FOOTER))


def _setup_styles(doc):
    """Configure Normal style with David 12pt RTL."""
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_PT)
    style.font.rtl = True

    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = etree.SubElement(style.element, qn("w:rPr"))
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = etree.SubElement(rPr, qn("w:szCs"))
    szCs.set(qn("w:val"), str(FONT_SIZE_HALF_POINTS))
    # RTL and language on style
    if rPr.find(qn("w:rtl")) is None:
        etree.SubElement(rPr, qn("w:rtl"))
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        lang = etree.SubElement(rPr, qn("w:lang"))
    lang.set(qn("w:bidi"), "he-IL")

    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = style.element.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        etree.SubElement(pPr, qn("w:bidi"))
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = etree.SubElement(pPr, qn("w:spacing"))
    sp.set(qn("w:before"), str(PARA_BEFORE))
    sp.set(qn("w:after"), str(PARA_AFTER))
    sp.set(qn("w:line"), str(LINE_SPACING))
    sp.set(qn("w:lineRule"), LINE_RULE)


def _setup_numbering(doc):
    """Create numbering definition per SKILL.md."""
    try:
        numbering_part = doc.part.numbering_part
    except Exception:
        dummy = doc.add_paragraph("", style="List Number")
        numbering_part = doc.part.numbering_part
        dummy._element.getparent().remove(dummy._element)

    numbering_elm = numbering_part.element

    abstract_num_xml = f"""
    <w:abstractNum w:abstractNumId="0" xmlns:w="{WNS}">
        <w:multiLevelType w:val="hybridMultilevel"/>
        <w:lvl w:ilvl="0">
            <w:start w:val="1"/>
            <w:numFmt w:val="decimal"/>
            <w:lvlText w:val="%1."/>
            <w:lvlJc w:val="left"/>
            <w:pPr>
                <w:ind w:left="360" w:hanging="360"/>
            </w:pPr>
            <w:rPr>
                <w:rFonts w:ascii="David" w:hAnsi="David" w:cs="David"/>
                <w:b w:val="0"/>
                <w:bCs w:val="0"/>
                <w:sz w:val="24"/>
                <w:szCs w:val="24"/>
                <w:lang w:bidi="he-IL"/>
            </w:rPr>
        </w:lvl>
    </w:abstractNum>
    """
    numbering_elm.insert(0, etree.fromstring(abstract_num_xml))

    num_xml = f"""
    <w:num w:numId="2" xmlns:w="{WNS}">
        <w:abstractNumId w:val="0"/>
    </w:num>
    """
    numbering_elm.append(etree.fromstring(num_xml))


# ══════════════════════════════════════════════════════════════════════════════
# PARAGRAPH HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _set_rtl_bidi(p):
    """Set bidi on paragraph pPr."""
    pPr = p._element.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        etree.SubElement(pPr, qn("w:bidi"))


def _set_run_font(run, size=12, bold=False, underline=False):
    """Configure run with David font, RTL, bidi language, complex-script sizes."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    run.font.rtl = True
    rPr = run._element.get_or_add_rPr()
    # rFonts with cs
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    # bCs for bold complex script
    if bold:
        if rPr.find(qn("w:bCs")) is None:
            etree.SubElement(rPr, qn("w:bCs"))
    # szCs
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = etree.SubElement(rPr, qn("w:szCs"))
    szCs.set(qn("w:val"), str(size * 2))
    # w:rtl
    if rPr.find(qn("w:rtl")) is None:
        etree.SubElement(rPr, qn("w:rtl"))
    # w:lang bidi="he-IL"
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        lang = etree.SubElement(rPr, qn("w:lang"))
    lang.set(qn("w:bidi"), "he-IL")


def _set_spacing(p):
    pPr = p._element.get_or_add_pPr()
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = etree.SubElement(pPr, qn("w:spacing"))
    sp.set(qn("w:before"), str(PARA_BEFORE))
    sp.set(qn("w:after"), str(PARA_AFTER))
    sp.set(qn("w:line"), str(LINE_SPACING))
    sp.set(qn("w:lineRule"), LINE_RULE)


def _add_numbering(p):
    pPr = p._element.get_or_add_pPr()
    numPr = etree.SubElement(pPr, qn("w:numPr"))
    ilvl = etree.SubElement(numPr, qn("w:ilvl"))
    ilvl.set(qn("w:val"), "0")
    numId_el = etree.SubElement(numPr, qn("w:numId"))
    numId_el.set(qn("w:val"), "2")
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = etree.SubElement(pPr, qn("w:ind"))
    ind.set(qn("w:left"), str(NUM_INDENT_LEFT))
    ind.set(qn("w:right"), str(NUM_INDENT_RIGHT))
    ind.set(qn("w:hanging"), str(NUM_INDENT_HANGING))


def _add_section_header(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_rtl_bidi(p)
    _set_spacing(p)
    pPr = p._element.get_or_add_pPr()
    ind = etree.SubElement(pPr, qn("w:ind"))
    ind.set(qn("w:left"), str(HDR_INDENT_LEFT))
    ind.set(qn("w:right"), str(HDR_INDENT_RIGHT))
    ind.set(qn("w:firstLine"), str(HDR_INDENT_FIRSTLINE))
    run = p.add_run(text)
    _set_run_font(run, bold=True, underline=True)
    return p


def _add_numbered_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_rtl_bidi(p)
    _set_spacing(p)
    _add_numbering(p)
    run = p.add_run(text)
    _set_run_font(run)
    return p


def _add_plain_para(doc, text, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12):
    p = doc.add_paragraph()
    p.alignment = alignment
    _set_rtl_bidi(p)
    _set_spacing(p)
    if text:
        run = p.add_run(text)
        _set_run_font(run, size=size, bold=bold)
    return p


def _add_appendix_ref(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_rtl_bidi(p)
    _set_spacing(p)
    pPr = p._element.get_or_add_pPr()
    ind = etree.SubElement(pPr, qn("w:ind"))
    ind.set(qn("w:left"), str(NUM_INDENT_LEFT))
    ind.set(qn("w:right"), str(NUM_INDENT_RIGHT))
    arrow_run = p.add_run("◄  ")
    _set_run_font(arrow_run, bold=True, underline=False)
    text_run = p.add_run(text)
    _set_run_font(text_run, bold=True, underline=True)
    return p


def _add_calculation_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_rtl_bidi(p)
    _set_spacing(p)
    pPr = p._element.get_or_add_pPr()
    ind = etree.SubElement(pPr, qn("w:ind"))
    ind.set(qn("w:left"), str(NUM_INDENT_LEFT))
    ind.set(qn("w:right"), str(NUM_INDENT_RIGHT))
    run = p.add_run(text)
    _set_run_font(run)
    return p


def _add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl_bidi(p)
    _set_spacing(p)
    run = p.add_run(text)
    _set_run_font(run, size=16, bold=True)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# TABLE HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _setup_table_bidi(table, width=TABLE_WIDTH):
    """Rebuild tblPr to ensure bidiVisual is BEFORE tblW.

    python-docx may create a default tblW. We remove it and rebuild
    the correct order: bidiVisual, then tblW.
    """
    tblEl = table._element
    tblPr = tblEl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = etree.SubElement(tblEl, qn("w:tblPr"))
        tblEl.insert(0, tblPr)

    # Remove any existing bidiVisual and tblW
    for tag in ["w:bidiVisual", "w:tblW"]:
        for existing in tblPr.findall(qn(tag)):
            tblPr.remove(existing)

    # Insert bidiVisual at position 0
    bidi = etree.SubElement(tblPr, qn("w:bidiVisual"))
    tblPr.insert(0, bidi)

    # Insert tblW right after bidiVisual
    tblW = etree.SubElement(tblPr, qn("w:tblW"))
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(width))
    tblPr.insert(1, tblW)

    return tblPr


def _setup_table_grid(table, col_widths):
    """Set up table grid columns, removing any defaults."""
    tblEl = table._element
    tblGrid = tblEl.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = etree.SubElement(tblEl, qn("w:tblGrid"))
    else:
        for gc in tblGrid.findall(qn("w:gridCol")):
            tblGrid.remove(gc)
    for w in col_widths:
        gc = etree.SubElement(tblGrid, qn("w:gridCol"))
        gc.set(qn("w:w"), str(w))


def _set_cell_rtl(cell, text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pPr = p._element.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        etree.SubElement(pPr, qn("w:bidi"))
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        pPr.remove(ind)
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = etree.SubElement(pPr, qn("w:spacing"))
    sp.set(qn("w:before"), "40")
    sp.set(qn("w:after"), "40")
    sp.set(qn("w:line"), "276")
    sp.set(qn("w:lineRule"), "auto")
    if text:
        for line_idx, line in enumerate(text.split("\n")):
            if line_idx > 0:
                run = p.add_run()
                run.add_break()
            run = p.add_run(line)
            _set_run_font(run, size=size, bold=bold)
    tc = cell._element
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(tc, qn("w:tcPr"))
        tc.insert(0, tcPr)


def _set_cell_multiline(cell, lines_spec):
    cell.text = ""
    for idx, (text, bold, size, alignment) in enumerate(lines_spec):
        if idx == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.alignment = alignment
        pPr = p._element.get_or_add_pPr()
        if pPr.find(qn("w:bidi")) is None:
            etree.SubElement(pPr, qn("w:bidi"))
        ind = pPr.find(qn("w:ind"))
        if ind is not None:
            pPr.remove(ind)
        sp = pPr.find(qn("w:spacing"))
        if sp is None:
            sp = etree.SubElement(pPr, qn("w:spacing"))
        sp.set(qn("w:before"), "20")
        sp.set(qn("w:after"), "20")
        sp.set(qn("w:line"), "240")
        sp.set(qn("w:lineRule"), "auto")
        if text:
            run = p.add_run(text)
            _set_run_font(run, size=size, bold=bold)


def _shade_cell(cell, fill_color, font_color=None):
    tc = cell._element
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(tc, qn("w:tcPr"))
        tc.insert(0, tcPr)
    shd = etree.SubElement(tcPr, qn("w:shd"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    if font_color:
        for run in cell.paragraphs[0].runs:
            rPr = run._element.get_or_add_rPr()
            color = rPr.find(qn("w:color"))
            if color is None:
                color = etree.SubElement(rPr, qn("w:color"))
            color.set(qn("w:val"), font_color)


def _add_table_borders(tblPr, style="single"):
    """Add visible borders to a table."""
    # Remove existing borders first
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBorders = etree.SubElement(tblPr, qn("w:tblBorders"))
    val = style if style != "none" else "none"
    sz = "4" if style != "none" else "0"
    color = "000000" if style != "none" else "auto"
    for bn in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = etree.SubElement(tblBorders, qn(f"w:{bn}"))
        b.set(qn("w:val"), val)
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)


def _set_cell_valign(cell, val="bottom"):
    tc = cell._element
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = etree.SubElement(tc, qn("w:tcPr"))
        tc.insert(0, tcPr)
    va = etree.SubElement(tcPr, qn("w:vAlign"))
    va.set(qn("w:val"), val)


def _add_summary_table(doc, claims_dict, total_amount):
    num_rows = len(claims_dict) + 2  # header + data + total
    tbl = doc.add_table(rows=num_rows, cols=2)

    # Setup bidiVisual BEFORE tblW
    tblPr = _setup_table_bidi(tbl)
    _setup_table_grid(tbl, [SUMMARY_COL_RIGHT, SUMMARY_COL_LEFT])
    _add_table_borders(tblPr, "single")

    # Header row
    _set_cell_rtl(tbl.rows[0].cells[0], "רכיב תביעה", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_rtl(tbl.rows[0].cells[1], "סכום (₪)", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _shade_cell(tbl.rows[0].cells[0], HEADER_BG, HEADER_FG)
    _shade_cell(tbl.rows[0].cells[1], HEADER_BG, HEADER_FG)

    # Data rows
    for i, (key, claim) in enumerate(claims_dict.items()):
        row_idx = i + 1
        _set_cell_rtl(tbl.rows[row_idx].cells[0], claim["name"], bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_rtl(tbl.rows[row_idx].cells[1], f"{claim['amount']:,.0f} ₪",
                      alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Total row (shaded)
    last_row = num_rows - 1
    _set_cell_rtl(tbl.rows[last_row].cells[0], 'סה"כ', bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    _set_cell_rtl(tbl.rows[last_row].cells[1], f"{total_amount:,.0f} ₪", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _shade_cell(tbl.rows[last_row].cells[0], TOTAL_ROW_BG)
    _shade_cell(tbl.rows[last_row].cells[1], TOTAL_ROW_BG)

    return tbl


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

def _build_cover_page(doc, form_data, claims, total):
    """Build the cover page: header table, parties table, summary table, title."""
    plaintiff_name = form_data.get("plaintiff_name", "")
    plaintiff_id = form_data.get("plaintiff_id", "")
    plaintiff_address = form_data.get("plaintiff_address", "")
    defendant_name = form_data.get("defendant_name", "")
    defendant_id = form_data.get("defendant_id", "")
    defendant_address = form_data.get("defendant_address", "")
    court_name = form_data.get("court_header", "בית הדין האזורי לעבודה בתל אביב")
    gender = form_data.get("gender", "male")
    pronoun = "התובע" if gender == "male" else "התובעת"
    defendant_label = "הנתבע" if form_data.get("defendant_type") == "individual" else "הנתבעת"
    attorney_name = form_data.get("attorney_name", "")
    attorney_id = form_data.get("attorney_id", "")
    firm_name = form_data.get("firm_name", "")
    firm_address = form_data.get("firm_address", "")
    firm_phone = form_data.get("firm_phone", "")
    firm_fax = form_data.get("firm_fax", "")
    firm_email = form_data.get("firm_email", "")

    # ══════════════════════════════════════════════════════════════════════
    # HEADER TABLE: 2 columns with bidiVisual
    # With bidiVisual: cells[0] = RIGHT side, cells[1] = LEFT side
    # RIGHT = סע"ש / בפני
    # LEFT  = court name
    # ══════════════════════════════════════════════════════════════════════
    hdr_tbl = doc.add_table(rows=1, cols=2)
    tblPr = _setup_table_bidi(hdr_tbl)
    _setup_table_grid(hdr_tbl, [4513, 4513])
    _add_table_borders(tblPr, "none")

    # Split court name into base + location
    court_base = court_name
    court_location = ""
    if " ב" in court_name:
        parts = court_name.rsplit(" ב", 1)
        if len(parts) == 2:
            court_base = parts[0]
            court_location = "ב" + parts[1]

    # cells[0] = RIGHT side: סע"ש and בפני
    _set_cell_multiline(hdr_tbl.rows[0].cells[0], [
        ('סע"ש ________', False, 11, WD_ALIGN_PARAGRAPH.RIGHT),
        ("בפני _________", False, 11, WD_ALIGN_PARAGRAPH.RIGHT),
    ])

    # cells[1] = LEFT side: court name
    court_lines = [(court_base, True, 12, WD_ALIGN_PARAGRAPH.LEFT)]
    if court_location:
        court_lines.append((court_location, True, 12, WD_ALIGN_PARAGRAPH.LEFT))
    _set_cell_multiline(hdr_tbl.rows[0].cells[1], court_lines)

    # ══════════════════════════════════════════════════════════════════════
    # PARTIES TABLE: 2 columns with bidiVisual
    # cells[0] = RIGHT (content), cells[1] = LEFT (label)
    # ══════════════════════════════════════════════════════════════════════
    parties_tbl = doc.add_table(rows=5, cols=2)
    pt_tblPr = _setup_table_bidi(parties_tbl)
    _setup_table_grid(parties_tbl, [7026, 2000])
    _add_table_borders(pt_tblPr, "single")

    # Row 0: "בעניין:"
    _set_cell_rtl(parties_tbl.rows[0].cells[0], "בעניין:", bold=True)
    _set_cell_rtl(parties_tbl.rows[0].cells[1], "")

    # Row 1: Plaintiff details (cells[0] RIGHT) | label (cells[1] LEFT)
    plaintiff_lines = []
    name_id = f"{plaintiff_name}, ת.ז. {plaintiff_id}" if plaintiff_id else plaintiff_name
    plaintiff_lines.append((name_id, True, 12, WD_ALIGN_PARAGRAPH.RIGHT))
    if plaintiff_address:
        plaintiff_lines.append((plaintiff_address, False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    if attorney_name:
        plaintiff_lines.append((f'באמצעות ב"כ עוה"ד {attorney_name}', False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    if firm_name:
        fn = f"ממשרד {firm_name}" if not firm_name.startswith("ממשרד") else firm_name
        plaintiff_lines.append((fn, False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    if firm_address:
        for part in firm_address.split(","):
            part = part.strip()
            if part:
                plaintiff_lines.append((part, False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    contact_parts = []
    if firm_phone:
        contact_parts.append(f"טל': {firm_phone}")
    if firm_fax:
        contact_parts.append(f"פקסי': {firm_fax}")
    if contact_parts:
        plaintiff_lines.append((" ".join(contact_parts), False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    if firm_email:
        plaintiff_lines.append((firm_email, False, 11, WD_ALIGN_PARAGRAPH.RIGHT))

    _set_cell_multiline(parties_tbl.rows[1].cells[0], plaintiff_lines)
    _set_cell_rtl(parties_tbl.rows[1].cells[1], pronoun, bold=True)
    _set_cell_valign(parties_tbl.rows[1].cells[1], "bottom")

    # Row 2: "- נגד -"
    _set_cell_rtl(parties_tbl.rows[2].cells[0], "- נגד -", bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_rtl(parties_tbl.rows[2].cells[1], "")

    # Row 3: Defendant
    defendant_lines = [(defendant_name, True, 12, WD_ALIGN_PARAGRAPH.RIGHT)]
    if defendant_id:
        defendant_lines.append((f"ח.פ {defendant_id}", False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    if defendant_address:
        defendant_lines.append((defendant_address, False, 11, WD_ALIGN_PARAGRAPH.RIGHT))
    _set_cell_multiline(parties_tbl.rows[3].cells[0], defendant_lines)
    _set_cell_rtl(parties_tbl.rows[3].cells[1], defendant_label, bold=True)
    _set_cell_valign(parties_tbl.rows[3].cells[1], "bottom")

    # Row 4: Claim nature and amount
    amount_str = f"{total:,.0f} ₪"
    _set_cell_multiline(parties_tbl.rows[4].cells[0], [
        ("מהות התביעה: הצהרתית וכספית", True, 11, WD_ALIGN_PARAGRAPH.RIGHT),
        (f"סכום התביעה: {amount_str}", True, 11, WD_ALIGN_PARAGRAPH.RIGHT),
    ])
    _set_cell_rtl(parties_tbl.rows[4].cells[1], "")

    # ── Summary table on cover page ──
    _add_plain_para(doc, "")
    _add_summary_table(doc, claims, total)

    # ── Title ──
    _add_title(doc, "כ ת ב    ת ב י ע ה")


# ══════════════════════════════════════════════════════════════════════════════
# SIGNATURE BLOCK
# ══════════════════════════════════════════════════════════════════════════════

def _add_signature_block(doc, form_data):
    gender = form_data.get("gender", "male")
    pronoun = "התובע" if gender == "male" else "התובעת"
    attorney_name = form_data.get("attorney_name", "")
    attorney_id = form_data.get("attorney_id", "")

    _add_plain_para(doc, "")

    sig_table = doc.add_table(rows=1, cols=2)
    sig_tblPr = _setup_table_bidi(sig_table)
    _setup_table_grid(sig_table, [SIG_COL_SPACER, SIG_COL_SIG])
    _add_table_borders(sig_tblPr, "none")

    # Empty spacer (cells[0] = right side with bidiVisual)
    _set_cell_rtl(sig_table.rows[0].cells[0], "")

    # Signature cell (cells[1] = left side)
    sig_cell = sig_table.rows[0].cells[1]
    if attorney_name and attorney_id:
        sig_text = f'{attorney_name}, עו"ד\nמ.ר. {attorney_id}\nב"כ {pronoun}'
    else:
        sig_text = f'__________________\nב"כ {pronoun}'
    _set_cell_rtl(sig_cell, sig_text, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Top border on signature cell (serves as signature line)
    sig_tc = sig_cell._element
    sig_tcPr = sig_tc.find(qn("w:tcPr"))
    if sig_tcPr is None:
        sig_tcPr = etree.SubElement(sig_tc, qn("w:tcPr"))
        sig_tc.insert(0, sig_tcPr)
    sig_tcBorders = etree.SubElement(sig_tcPr, qn("w:tcBorders"))
    top_border = etree.SubElement(sig_tcBorders, qn("w:top"))
    top_border.set(qn("w:val"), "single")
    top_border.set(qn("w:sz"), "4")
    top_border.set(qn("w:space"), "0")
    top_border.set(qn("w:color"), "auto")


# ══════════════════════════════════════════════════════════════════════════════
# TEXT PARSING & GENDER
# ══════════════════════════════════════════════════════════════════════════════

def _parse_sections(raw_text):
    """Parse plain text with === TITLE === delimiters into sections."""
    sections = []
    current_title = ""
    current_lines = []

    for line in raw_text.split("\n"):
        stripped = line.strip()
        match = re.match(r"^===\s*(.+?)\s*===$", stripped)
        if match:
            if current_title or current_lines:
                sections.append({
                    "title": current_title,
                    "lines": [l for l in current_lines if l.strip()],
                })
            current_title = match.group(1)
            current_lines = []
        else:
            current_lines.append(stripped)

    if current_title or current_lines:
        sections.append({
            "title": current_title,
            "lines": [l for l in current_lines if l.strip()],
        })

    return sections


def _clean_line(text):
    if not text:
        return text
    # Remove AI-added numbering prefixes
    text = re.sub(r"^\d+\.\s+", "", text)
    text = re.sub(r"  +", " ", text)
    return text.strip()


def _has_hebrew(text):
    return bool(re.search(r"[\u0590-\u05FF]", text))


def _apply_gender_to_doc(doc, gender):
    """Apply gender fixes to all paragraphs in the document body."""
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text:
                run.text = fix_gender(run.text, gender)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text:
                            run.text = fix_gender(run.text, gender)


def _ensure_hebrew_proofing(doc):
    """Final pass: ensure EVERY run in the document has w:rtl, w:rFonts cs=David, w:lang bidi=he-IL.
    Also ensure every paragraph has w:bidi in pPr."""

    def _proof_paragraph(p_element):
        pPr = p_element.find(qn("w:pPr"))
        if pPr is None:
            pPr = etree.SubElement(p_element, qn("w:pPr"))
            p_element.insert(0, pPr)
        if pPr.find(qn("w:bidi")) is None:
            etree.SubElement(pPr, qn("w:bidi"))

    def _proof_run(r_element):
        rPr = r_element.find(qn("w:rPr"))
        if rPr is None:
            rPr = etree.SubElement(r_element, qn("w:rPr"))
            r_element.insert(0, rPr)
        # w:rtl
        if rPr.find(qn("w:rtl")) is None:
            etree.SubElement(rPr, qn("w:rtl"))
        # w:rFonts cs=David
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        if not rFonts.get(qn("w:cs")):
            rFonts.set(qn("w:cs"), FONT_NAME)
        # w:lang bidi=he-IL
        lang = rPr.find(qn("w:lang"))
        if lang is None:
            lang = etree.SubElement(rPr, qn("w:lang"))
        if not lang.get(qn("w:bidi")):
            lang.set(qn("w:bidi"), "he-IL")

    # Process all paragraphs in document body
    body = doc.element.body
    for p in body.iter(qn("w:p")):
        _proof_paragraph(p)
        for r in p.iter(qn("w:r")):
            _proof_run(r)

    # Process all tables (including nested)
    for tbl in body.iter(qn("w:tbl")):
        for p in tbl.iter(qn("w:p")):
            _proof_paragraph(p)
            for r in p.iter(qn("w:r")):
                _proof_run(r)
